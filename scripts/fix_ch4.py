"""
Transforms Ch4_MoHinhHoaCauTruc.docx:
1. Renumbers all 3.x / 3.x.y headings to 4.x / 4.x.y
2. Fixes body text references: "Chương III" -> "Chương IV"
3. Updates intro order sentence
4. Per-UC: reorders subsections so Object Diagram comes FIRST,
   then Class Diagram, then Scenario (was: class, scenario, object)
   and renumbers .1/.2/.3 accordingly
"""
import copy
import re
from docx import Document
from docx.oxml.ns import qn

DOC_IN  = r"Project/Ch4_MoHinhHoaCauTruc.docx"
DOC_OUT = r"Project/Ch4_MoHinhHoaCauTruc.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_para_text_preserve_fmt(para, new_text):
    """Replace the full text of a paragraph while keeping the first run's formatting."""
    # Clear all runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def full_text(para):
    return "".join(r.text for r in para.runs) or para.text


def replace_text_in_para(para, old, new):
    """Replace text across runs if needed, or in each run individually."""
    if old in full_text(para):
        # Try simple per-run replacement first
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        # Fallback: rebuild as single run
        combined = full_text(para)
        if old in combined:
            set_para_text_preserve_fmt(para, combined.replace(old, new))
            return True
    return False


def swap_paragraphs(doc, idx_a, idx_b):
    """Swap the XML elements of two paragraphs by index."""
    paras = doc.paragraphs
    elem_a = paras[idx_a]._element
    elem_b = paras[idx_b]._element
    parent = elem_a.getparent()
    pos_a  = list(parent).index(elem_a)
    pos_b  = list(parent).index(elem_b)
    # Remove both, reinsert in swapped positions
    parent.remove(elem_a)
    parent.remove(elem_b)
    # Insert at lower index first
    if pos_a < pos_b:
        parent.insert(pos_a, elem_b)
        parent.insert(pos_b, elem_a)
    else:
        parent.insert(pos_b, elem_a)
        parent.insert(pos_a, elem_b)


# ── load ──────────────────────────────────────────────────────────────────────

doc = Document(DOC_IN)
paras = doc.paragraphs

# ── 1. Fix "Chương III" → "Chương IV" in body text ──────────────────────────
for p in paras:
    replace_text_in_para(p, "Chương III", "Chương IV")

# ── 2. Renumber section headings 3.x → 4.x ──────────────────────────────────
heading_re = re.compile(r'^(3\.\d+\.?\d*)\.')

for p in paras:
    if p.style.name.startswith("Heading"):
        t = full_text(p)
        m = heading_re.match(t)
        if m:
            old_num = m.group(1)
            new_num = "4" + old_num[1:]   # replace leading "3" with "4"
            new_text = new_num + "." + t[len(old_num)+1:]
            set_para_text_preserve_fmt(p, new_text)

# ── 3. Update intro-order sentence ───────────────────────────────────────────
OLD_ORDER = (
    "(1) sơ đồ lớp lĩnh vực, (2) kịch bản sử dụng cụ thể, "
    "(3) sơ đồ đối tượng minh họa, (4) thẻ CRC cho các lớp chính."
)
NEW_ORDER = (
    "(1) sơ đồ đối tượng minh họa (kịch bản cụ thể), "
    "(2) sơ đồ lớp lĩnh vực, (3) kịch bản sử dụng, "
    "(4) thẻ CRC cho các lớp chính."
)
for p in paras:
    replace_text_in_para(p, OLD_ORDER, NEW_ORDER)

# ── 4. Per-UC: reorder subsections ───────────────────────────────────────────
# After step 2 renaming, the headings now look like:
#   4.1. UC-01 ...   (Heading 2)
#   4.1.1. Sơ đồ lớp lĩnh vực   -> will become 4.1.2
#   [content]
#   4.1.2. Kịch bản sử dụng     -> will become 4.1.3
#   [content]
#   4.1.3. Sơ đồ đối tượng      -> will become 4.1.1
#   [content]
#   4.1.4. Thẻ CRC               -> stays 4.1.4
#
# We need to:
#   a) Physically reorder: move "Sơ đồ đối tượng" block BEFORE "Sơ đồ lớp lĩnh vực" block
#   b) Renumber headings: old .1→.2, old .2→.3, old .3→.1

# Because swap is easier done on the XML tree, we'll do block moves.
# Strategy: for each UC section, collect the paragraph indices of each sub-section
# and physically reorder them in the XML.

def find_uc_blocks(doc):
    """
    Returns list of dicts for each UC, each containing:
      'h2': index of the UC Heading 2
      'class': (heading_idx, content_idx)   -- Sơ đồ lớp lĩnh vực block
      'scenario': (heading_idx, content_idx) -- Kịch bản sử dụng block
      'object': (heading_idx, content_idx)   -- Sơ đồ đối tượng block
      'crc': (heading_idx, content_idx)      -- Thẻ CRC block
    """
    uc_blocks = []
    paras = doc.paragraphs
    n = len(paras)
    i = 0
    while i < n:
        p = paras[i]
        if p.style.name == "Heading 2" and re.match(r'^4\.\d+\.', full_text(p)):
            block = {"h2": i, "class": None, "scenario": None, "object": None, "crc": None}
            j = i + 1
            while j < n and not (paras[j].style.name == "Heading 2" and re.match(r'^4\.\d+\.', full_text(paras[j]))):
                t = full_text(paras[j])
                if paras[j].style.name == "Heading 3":
                    if "lớp lĩnh vực" in t:
                        block["class"] = (j, j+1)
                    elif "Kịch bản" in t or "kịch bản" in t:
                        block["scenario"] = (j, j+1)
                    elif "đối tượng" in t or "Đối tượng" in t:
                        block["object"] = (j, j+1)
                    elif "CRC" in t:
                        block["crc"] = (j, j+1)
                j += 1
            uc_blocks.append(block)
            i = j
        else:
            i += 1
    return uc_blocks


def move_element_before(parent, elem_to_move, reference_elem):
    """Move elem_to_move to be immediately before reference_elem in parent."""
    parent.remove(elem_to_move)
    ref_pos = list(parent).index(reference_elem)
    parent.insert(ref_pos, elem_to_move)


def reorder_uc_subsections(doc):
    """
    For each UC block, move the Sơ đồ đối tượng (+ its content paragraph)
    to come before Sơ đồ lớp lĩnh vực (+ its content paragraph).
    Then renumber: old .1→.2, old .2→.3, old .3→.1
    """
    uc_re = re.compile(r'^(4\.\d+)\.')
    sub_re = re.compile(r'^(4\.\d+\.)(\d+)\.')

    uc_blocks = find_uc_blocks(doc)
    paras = doc.paragraphs

    for block in uc_blocks:
        if not all([block["class"], block["scenario"], block["object"]]):
            print(f"  WARNING: incomplete block for UC at h2={block['h2']}, skipping reorder")
            continue

        class_h_idx,    class_c_idx    = block["class"]
        scenario_h_idx, scenario_c_idx = block["scenario"]
        object_h_idx,   object_c_idx   = block["object"]

        body = paras[class_h_idx]._element.getparent()

        # We need to move [object_heading, object_content] to before [class_heading]
        obj_h_elem = paras[object_h_idx]._element
        obj_c_elem = paras[object_c_idx]._element
        cls_h_elem = paras[class_h_idx]._element

        # Move object content first, then heading (so heading ends up right before class)
        move_element_before(body, obj_c_elem, cls_h_elem)
        move_element_before(body, obj_h_elem, cls_h_elem)  # obj_h will be before obj_c and before cls_h

    # After physical reordering, re-read paragraphs and renumber
    # Mapping: within each 4.x. section -> old sub 1->2, 2->3, 3->1, 4 stays
    # We use a two-pass rename to avoid conflicts: first rename to temps, then to final

    # Pass 1: rename to temporary values
    # old .1 → temp .11, old .2 → temp .22, old .3 → temp .33
    temp_map = {"1": "11", "2": "22", "3": "33"}
    for p in doc.paragraphs:
        if p.style.name == "Heading 3":
            t = full_text(p)
            m = sub_re.match(t)
            if m and m.group(2) in temp_map:
                new_t = m.group(1) + temp_map[m.group(2)] + "." + t[len(m.group(0)):]
                set_para_text_preserve_fmt(p, new_t)

    # Pass 2: temp → final  (11→2, 22→3, 33→1)
    temp_re = re.compile(r'^(4\.\d+\.)(11|22|33)\.')
    final_map = {"11": "2", "22": "3", "33": "1"}
    for p in doc.paragraphs:
        if p.style.name == "Heading 3":
            t = full_text(p)
            m = temp_re.match(t)
            if m and m.group(2) in final_map:
                new_t = m.group(1) + final_map[m.group(2)] + "." + t[len(m.group(0)):]
                set_para_text_preserve_fmt(p, new_t)


reorder_uc_subsections(doc)

# ── 5. Save ──────────────────────────────────────────────────────────────────
doc.save(DOC_OUT)
print("Done. Saved to", DOC_OUT)
