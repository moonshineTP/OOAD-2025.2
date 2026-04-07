"""
scripts/fill_ssd_crc.py
Fills in sections 4.3.4 (SSD) and 4.3.5 (CRC cards) for UC-03.

4.3.4:
  - Intro paragraph
  - SSD message contract table (5 outgoing + 5 return messages)
  - PlantUML code block as a pre-formatted paragraph
  - Drawio placeholder line

4.3.5:
  - Keeps the existing template table as-is (as a reference example)
  - Inserts 5 filled CRC card tables after the "[Chèn thẻ CRC]" placeholder:
    ChuyenDi, LoTrinh, ViTriGPS, KienHang, BangChungGiaoHang
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document(DOC_PATH)
body = doc.element.body

# ── helpers ───────────────────────────────────────────────────────────────────

def gettxt(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{NS}}}t"))

def detach(obj):
    e = obj._element if hasattr(obj, "_element") else obj
    e.getparent().remove(e)
    return e

def make_para(text="", style="Normal", bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return detach(p)

def make_mixed(segments):
    """segments = [(text, bold, italic)]"""
    p = doc.add_paragraph(style="Normal")
    for text, bold, italic in segments:
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
    return detach(p)

def _make_border(tag, val="single", sz="6", color="000000", space="0"):
    el = OxmlElement(tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:color"), color)
    el.set(qn("w:space"), space)
    return el

def apply_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    else:
        for child in list(tcBorders):
            tcBorders.remove(child)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tcBorders.append(_make_border(f"w:{side}"))

def apply_table_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            apply_cell_borders(cell)

def make_table(headers, rows, header_bold=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = "TableNormal"
    except Exception:
        pass
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if header_bold:
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
    apply_table_borders(tbl)
    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def make_crc_table(tbl_name, tbl_id, layer, desc, attrs, responsibilities, partners):
    """Build a 7-row × 2-col CRC vertical table matching the template."""
    tbl = doc.add_table(rows=7, cols=2)
    try:
        tbl.style = "TableNormal"
    except Exception:
        pass
    data = [
        ("Tên lớp",         tbl_name),
        ("ID",               tbl_id),
        ("Tầng kiến trúc",  layer),
        ("Mô tả",           desc),
        ("Các thuộc tính",  attrs),
        ("Các trách nhiệm", responsibilities),
        ("Các đối tác",     partners),
    ]
    for ri, (label, value) in enumerate(data):
        cells = tbl.rows[ri].cells
        cells[0].text = label
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
        cells[1].text = value
    apply_table_borders(tbl)
    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def make_code_block(lines):
    """Render as indented Normal paragraphs (Courier-style in Word via direct font)."""
    elems = []
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = None  # keep default size
        elems.append(detach(p))
    return elems

def insert_before(elems, ref_elem):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4.3.4 — SSD
# ═══════════════════════════════════════════════════════════════════════════════

children = list(body)
ssd_placeholder = None
for child in children:
    if "Chèn sơ đồ tuần tự hệ thống" in gettxt(child) and "4.3" in gettxt(list(body)[list(body).index(child)-2]):
        ssd_placeholder = child
        break
# More robust: find the one after the 4.3.4 heading
for i, child in enumerate(children):
    if "4.3.4" in gettxt(child) and "SSD" in gettxt(child):
        ssd_placeholder = children[i + 1]
        break

assert ssd_placeholder is not None, "SSD placeholder not found"

ssd_elems = []

# Intro
ssd_elems.append(make_para(
    "Sơ đồ tuần tự hệ thống (SSD) UC-03 mô tả chuỗi thông điệp trao đổi giữa "
    "Tài xế giao hàng (actor) và :Hệ thống LogiFast trong luồng sự kiện chính. "
    "Mỗi thông điệp gửi đi tương ứng với một bước trong đặc tả UC-03. "
    "SSD chỉ hiển thị hai đường sống: actor và :Hệ thống — không có Controller, "
    "DAO hay lớp nội bộ nào để đảm bảo đúng mức phân tích."
))

# Message contract table
ssd_elems.append(make_mixed([("Bảng thông điệp SSD — UC-03:", True, False)]))
ssd_elems.append(make_para(
    "Các thông điệp dưới đây được đặt tên theo quy ước camelCase tiếng Việt; "
    "tham số trong ngoặc là dữ liệu tối thiểu hệ thống cần nhận để xử lý bước tương ứng."
))

msg_h = ["#", "Hướng", "Tên thông điệp (tham số)", "Bước UC-03", "Ghi chú"]
msg_r = [
    ["1", "→ :Hệ thống", "xacNhanBatDauCa(maTaiXe)",
     "Bước 1 — bắt đầu ca", "Tạo ChuyenDi"],
    ["↩", ":Hệ thống →", "xacNhanChuyenDi(maChuyenDi, danhSachDon)",
     "—", "Trả về ChuyenDi + danh sách đơn cần giao"],
    ["2", "→ :Hệ thống", "quetMaQR(maQR)",
     "Bước 2 — xác nhận kiện hàng", "Xác định KienHang theo mã QR"],
    ["↩", ":Hệ thống →", "xacNhanKienHang(tenSanPham, diaChiGiao)",
     "—", "Trả về chi tiết đơn nếu QR hợp lệ"],
    ["3", "→ :Hệ thống", "batDauGiaoHang(maDonHang)",
     "Bước 3 — xuất phát", "DonHang.trangThai → Đang giao"],
    ["↩", ":Hệ thống →", "hienThiLoTrinh(maLoTrinh, khoangCach, thoiGianDuKien)",
     "—", "Tính LoTrinh tối ưu, trả ETA"],
    ["4", "→ :Hệ thống", "capNhatViTriGPS(viDo, kinhDo)",
     "Bước 4 — tracking hành trình", "Ghi ViTriGPS, loop mỗi 30 giây"],
    ["↩", ":Hệ thống →", "ghiNhanThanhCong()",
     "—", "Xác nhận đã lưu tọa độ"],
    ["5", "→ :Hệ thống", "xacNhanBanGiao(maDonHang, tenNguoiNhan, hinhAnh)",
     "Bước 6 — bàn giao kiện hàng", "Tạo BangChungGiaoHang"],
    ["↩", ":Hệ thống →", "xacNhanGiaoHangThanhCong(maBangChung)",
     "—", "DonHang.trangThai → Giao thành công; phát DeliveryCompleted"],
]
t, s = make_table(msg_h, msg_r)
ssd_elems.extend([t, s])

# PlantUML code
ssd_elems.append(make_mixed([("Lược đồ SSD (PlantUML):", True, False)]))
ssd_elems.append(make_para(
    "Đoạn mã PlantUML dưới đây mô tả đầy đủ chuỗi thông điệp; "
    "sơ đồ được render và chèn vào tệp "
    "Project/Diagrams/UC03_SSD.puml."
))
puml_lines = [
    "@startuml",
    "skinparam style strictuml",
    "",
    'actor "Tài xế giao hàng" as TX',
    'participant ":Hệ thống LogiFast" as SYS',
    "",
    "== Bước 1: Bắt đầu ca làm việc ==",
    "TX -> SYS : xacNhanBatDauCa(maTaiXe)",
    "SYS --> TX : xacNhanChuyenDi(maChuyenDi, danhSachDon)",
    "",
    "== Bước 2: Xác nhận lấy hàng tại kho ==",
    "TX -> SYS : quetMaQR(maQR)",
    "SYS --> TX : xacNhanKienHang(tenSanPham, diaChiGiao)",
    "",
    "== Bước 3: Bắt đầu giao hàng ==",
    "TX -> SYS : batDauGiaoHang(maDonHang)",
    "SYS --> TX : hienThiLoTrinh(maLoTrinh, khoangCach, thoiGianDuKien)",
    "",
    "== Bước 4: Tracking hành trình (lặp mỗi 30 giây) ==",
    "loop mỗi 30 giây",
    "  TX -> SYS : capNhatViTriGPS(viDo, kinhDo)",
    "  SYS --> TX : ghiNhanThanhCong()",
    "end",
    "",
    "== Bước 5–6: Bàn giao và tạo bằng chứng ==",
    "TX -> SYS : xacNhanBanGiao(maDonHang, tenNguoiNhan, hinhAnh)",
    "SYS --> TX : xacNhanGiaoHangThanhCong(maBangChung)",
    "@enduml",
]
ssd_elems.extend(make_code_block(puml_lines))

# Drawio placeholder
ssd_elems.append(make_para())
ssd_elems.append(make_mixed([
    ("[Chèn sơ đồ SSD UC-03 — file: ", False, False),
    ("Project/Diagrams/UC03_SSD.puml", True, False),
    ("]", False, False),
]))

# Replace placeholder
insert_before(ssd_elems, ssd_placeholder)
ssd_placeholder.getparent().remove(ssd_placeholder)
print(f"4.3.4: inserted {len(ssd_elems)} elements")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4.3.5 — CRC Cards
# ═══════════════════════════════════════════════════════════════════════════════

children = list(body)
crc_placeholder = None
for child in children:
    if "Chèn thẻ CRC các lớp chính" in gettxt(child):
        crc_placeholder = child
        break

assert crc_placeholder is not None, "CRC placeholder not found"

crc_elems = []

# Brief intro
crc_elems.append(make_para(
    "Thẻ CRC (Class–Responsibility–Collaborator) được lập cho 5 lớp lĩnh vực "
    "trung tâm của UC-03. Các lớp TaiXe và DonHang được lập thẻ CRC ở UC-02 và UC-01 "
    "tương ứng vì chúng xuất hiện lần đầu ở đó. Thuộc tính chỉ ghi tên — không có "
    "kiểu dữ liệu (quy tắc phân tích I.12)."
))
crc_elems.append(make_para())

# 5 CRC cards
crc_data = [
    (
        "ChuyenDi", "CRC-UC03-01", "Lĩnh vực (Domain)",
        "Đại diện cho một ca làm việc của tài xế, bao gồm một hoặc nhiều lộ trình giao hàng trong một ngày.",
        "maChuyenDi, thoiGianBatDau, thoiGianKetThuc",
        "1. Ghi nhận thời điểm bắt đầu và kết thúc ca làm việc.\n"
        "2. Liên kết tài xế với tập hợp các lộ trình trong ca.\n"
        "3. Cung cấp ngữ cảnh ca làm việc để tổng hợp hiệu suất.",
        "TaiXe, LoTrinh",
    ),
    (
        "LoTrinh", "CRC-UC03-02", "Lĩnh vực (Domain)",
        "Tuyến đường từ kho xuất phát đến điểm giao hàng cho một đơn cụ thể; "
        "lưu trữ chuỗi tọa độ GPS toàn hành trình.",
        "maLoTrinh, diemXuatPhat, diemDen, khoangCach, /thoiGianDuKien, trangThai",
        "1. Lưu thông tin tuyến đường (điểm đầu, điểm cuối, khoảng cách).\n"
        "2. Tập hợp các bản ghi ViTriGPS trong suốt hành trình.\n"
        "3. Theo dõi trạng thái: Đang giao → Hoàn tất.\n"
        "4. Cung cấp dữ liệu ETA cho ứng dụng tài xế.",
        "ChuyenDi, TaiXe, DonHang, ViTriGPS",
    ),
    (
        "ViTriGPS", "CRC-UC03-03", "Lĩnh vực (Domain)",
        "Một bản ghi tọa độ GPS tại một thời điểm cụ thể trong hành trình; "
        "được ghi nhận mỗi 30 giây.",
        "maViTri, viDo, kinhDo, thoiGian",
        "1. Lưu tọa độ địa lý (vĩ độ, kinh độ) tại thời điểm ghi nhận.\n"
        "2. Cung cấp dữ liệu cho tính năng theo dõi vị trí thời gian thực.\n"
        "3. Tạo ra vết hành trình (breadcrumb trail) để kiểm tra sau giao.",
        "LoTrinh",
    ),
    (
        "KienHang", "CRC-UC03-04", "Lĩnh vực (Domain)",
        "Đơn vị vật lý được vận chuyển; được xác nhận bằng mã QR khi tài xế lấy "
        "hàng tại kho.",
        "maKienHang, maQR, khoiLuong, trangThai",
        "1. Xác minh danh tính kiện hàng qua quét mã QR.\n"
        "2. Theo dõi trạng thái vật lý: Đang vận chuyển → Đã giao.\n"
        "3. Cung cấp thông tin khoiLuong để tính phí vận chuyển (UC-05).",
        "DonHang",
    ),
    (
        "BangChungGiaoHang", "CRC-UC03-05", "Lĩnh vực (Domain)",
        "Bằng chứng xác thực việc giao hàng thành công, lưu ảnh chụp biên nhận, "
        "tên người nhận và tọa độ GPS tại điểm giao.",
        "maBangChung, tenNguoiNhan, hinhAnh, toaDo, thoiGianGiao",
        "1. Lưu trữ ảnh biên nhận giao hàng (hinhAnh).\n"
        "2. Ghi nhận tên người nhận thực tế và vị trí giao hàng.\n"
        "3. Cung cấp bằng chứng cho quá trình xác nhận giao hàng ở UC-04.",
        "DonHang, XacNhanGiaoHang",
    ),
]

for args in crc_data:
    t, s = make_crc_table(*args)
    crc_elems.extend([t, s])

# Insert CRC cards before placeholder, then remove placeholder
insert_before(crc_elems, crc_placeholder)
crc_placeholder.getparent().remove(crc_placeholder)
print(f"4.3.5: inserted {len(crc_elems)} elements")


doc.save(DOC_PATH)
print(f"Done. Saved to {DOC_PATH}")
