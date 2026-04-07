"""
Transforms Ch4_MoHinhHoaCauTruc.docx (v2 — full clean rebuild of UC sections).

Verified order (from docs/structural-modeling-prep.md line 107 and grading rubric I.9, I.17):
  1. Kịch bản sử dụng     (.x.1)
  2. Sơ đồ đối tượng      (.x.2)
  3. Sơ đồ lớp lĩnh vực   (.x.3)
  4. Sơ đồ tuần tự (SSD)  (.x.4)  ← NEW (required by rubric I.17)
  5. Thẻ CRC              (.x.5)

Also fixes:
  • Orphaned object-diagram content paragraph (bug from v1 swap)
  • Intro sentence order
"""
import copy, re
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def ft(para):
    return "".join(r.text for r in para.runs) or para.text

def set_text(para, text):
    """Replace full text of paragraph, keeping first run's format."""
    for r in para.runs:
        r.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def clone_para(para):
    """Deep-copy a paragraph element."""
    return copy.deepcopy(para._element)

def new_heading3_elem(doc, text):
    """Create a new Heading 3 paragraph element with given text."""
    p = doc.add_paragraph(text, style="Heading 3")
    elem = p._element
    # detach from document body (we only want the element, not the appended para)
    elem.getparent().remove(elem)
    return elem

def new_normal_elem(doc, text):
    """Create a new Normal paragraph element with given text."""
    p = doc.add_paragraph(text, style="Normal")
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def insert_after(parent, new_elem, ref_elem):
    """Insert new_elem into parent immediately after ref_elem."""
    idx = list(parent).index(ref_elem)
    parent.insert(idx + 1, new_elem)

# ── load document ─────────────────────────────────────────────────────────────

doc = Document(DOC_PATH)

# ── fix intro sentence ────────────────────────────────────────────────────────

OLD_ORDERS = [
    # previous v1 result
    "(1) sơ đồ đối tượng minh họa (kịch bản cụ thể), (2) sơ đồ lớp lĩnh vực, (3) kịch bản sử dụng, (4) thẻ CRC cho các lớp chính.",
    # original
    "(1) sơ đồ lớp lĩnh vực, (2) kịch bản sử dụng cụ thể, (3) sơ đồ đối tượng minh họa, (4) thẻ CRC cho các lớp chính.",
]
NEW_ORDER = (
    "(1) kịch bản sử dụng cụ thể, (2) sơ đồ đối tượng minh họa, "
    "(3) sơ đồ lớp lĩnh vực, (4) sơ đồ tuần tự hệ thống (SSD), "
    "(5) thẻ CRC cho các lớp chính."
)

for para in doc.paragraphs:
    t = ft(para)
    for old in OLD_ORDERS:
        if old in t:
            new_t = t.replace(old, NEW_ORDER)
            set_text(para, new_t)
            break

# ── per-UC section rebuild ────────────────────────────────────────────────────
#
# Strategy:
#  1. Identify each UC block by its Heading 2 ("4.x.")
#  2. Collect all body paragraphs until next Heading 2 (or end)
#  3. From those, classify each paragraph by its content keywords
#  4. Remove all body paragraphs from the XML
#  5. Re-insert in correct order, creating new SSD placeholders
#  6. Renumber headings .1 to .5

UC_H2_RE = re.compile(r"^4\.\d+\.")

def classify(para):
    """Return one of: 'scenario_h', 'scenario_c', 'object_h', 'object_c',
                      'class_h', 'class_c', 'crc_h', 'crc_c', 'unknown'"""
    t = ft(para)
    s = para.style.name
    if s.startswith("Heading 3"):
        if "Kịch bản" in t or "kịch bản" in t:
            return "scenario_h"
        if "đối tượng" in t.lower():
            return "object_h"
        if "lớp lĩnh vực" in t:
            return "class_h"
        if "CRC" in t:
            return "crc_h"
    else:  # Normal or other body
        if "kịch bản cụ thể" in t or "bảng đối tượng" in t:
            return "scenario_c"
        if "sơ đồ đối tượng" in t.lower():
            return "object_c"
        if "sơ đồ lớp lĩnh vực" in t.lower():
            return "class_c"
        if "thẻ CRC" in t or "CRC" in t:
            return "crc_c"
    return "unknown"

paras = doc.paragraphs
body = doc.element.body

# Find UC block boundaries
uc_starts = []  # list of (h2_idx, next_boundary_idx)
for i, p in enumerate(paras):
    if p.style.name == "Heading 2" and UC_H2_RE.match(ft(p)):
        uc_starts.append(i)

uc_ranges = []
for k, start in enumerate(uc_starts):
    end = uc_starts[k + 1] if k + 1 < len(uc_starts) else len(paras)
    uc_ranges.append((start, end))

for h2_idx, end_idx in uc_ranges:
    h2_para = paras[h2_idx]
    h2_text = ft(h2_para)
    # Extract UC number e.g. "4.1"
    uc_num = re.match(r"^(4\.\d+)\.", h2_text).group(1)

    # Collect body paragraphs (everything after H2, before next H2)
    body_paras = paras[h2_idx + 1 : end_idx]

    # Classify them
    buckets = {
        "scenario_h": None, "scenario_c": None,
        "object_h": None,   "object_c": None,
        "class_h": None,    "class_c": None,
        "crc_h": None,      "crc_c": None,
    }
    for bp in body_paras:
        cls = classify(bp)
        if cls in buckets:
            buckets[cls] = bp._element

    # Warn if anything's missing
    for key, val in buckets.items():
        if val is None:
            print(f"  WARNING {uc_num}: missing '{key}'")

    # Remove all body paragraphs from DOM
    for bp in body_paras:
        bp._element.getparent().remove(bp._element)

    # Build new elements in target order:
    # .1 Kịch bản, .2 Đối tượng, .3 Lớp, .4 SSD (new), .5 CRC
    ordered_elems = []

    def h3(text):
        return new_heading3_elem(doc, text)

    def norm(text):
        return new_normal_elem(doc, text)

    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def set_elem_text(elem, text):
        """Set the full text of an lxml paragraph element."""
        runs = elem.findall(f".//{{{NS}}}r")
        if runs:
            for r in runs:
                for tn in r.findall(f"{{{NS}}}t"):
                    tn.text = ""
            first_t = runs[0].find(f"{{{NS}}}t")
            if first_t is not None:
                first_t.text = text
                first_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        else:
            r_elem = OxmlElement("w:r")
            t_elem = OxmlElement("w:t")
            t_elem.text = text
            r_elem.append(t_elem)
            elem.append(r_elem)

    def get_or_new_h3(key, fallback_text):
        elem = buckets[key]
        if elem is None:
            elem = h3(fallback_text)
        set_elem_text(elem, fallback_text)
        return elem

    def get_or_new_norm(key, fallback_text):
        elem = buckets[key]
        if elem is None:
            elem = norm(fallback_text)
        return elem

    # 1. Kịch bản sử dụng
    sc_h = get_or_new_h3("scenario_h", f"{uc_num}.1. Kịch bản sử dụng")
    sc_c = get_or_new_norm("scenario_c", "[Mô tả kịch bản cụ thể và bảng đối tượng]")
    ordered_elems.extend([sc_h, sc_c])

    # 2. Sơ đồ đối tượng
    obj_h = get_or_new_h3("object_h", f"{uc_num}.2. Sơ đồ đối tượng")
    obj_c = get_or_new_norm("object_c", "[Chèn sơ đồ đối tượng]")
    ordered_elems.extend([obj_h, obj_c])

    # 3. Sơ đồ lớp lĩnh vực
    cls_h = get_or_new_h3("class_h", f"{uc_num}.3. Sơ đồ lớp lĩnh vực")
    cls_c = get_or_new_norm("class_c", "[Chèn sơ đồ lớp lĩnh vực]")
    ordered_elems.extend([cls_h, cls_c])

    # 4. Sơ đồ tuần tự hệ thống (SSD) — NEW
    ssd_h = h3(f"{uc_num}.4. Sơ đồ tuần tự hệ thống (SSD)")
    ssd_c = norm("[Chèn sơ đồ tuần tự hệ thống (SSD)]")
    ordered_elems.extend([ssd_h, ssd_c])

    # 5. Thẻ CRC
    crc_h = get_or_new_h3("crc_h", f"{uc_num}.5. Thẻ CRC")
    crc_c = get_or_new_norm("crc_c", "[Chèn thẻ CRC các lớp chính]")
    ordered_elems.extend([crc_h, crc_c])

    # Insert all after the H2 element
    ref = h2_para._element
    for elem in reversed(ordered_elems):
        insert_after(body, elem, ref)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Done.")
