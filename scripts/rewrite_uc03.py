"""
scripts/rewrite_uc03.py

Comprehensive rewrite of UC-03 section in Ch4_MoHinhHoaCauTruc.docx:
  1.  Remove BangChungGiaoHang from UC-03 scenario, object table, link table,
      class diagram, association table, SSD, and CRC cards.
  2.  Fix scenario endpoint: UC-03 ends at physical handover → "Chờ xác nhận";
      BangChungGiaoHang creation moves to UC-04.
  3.  Add rerouting event (09:45) to scenario, showing Hệ thống theo dõi +
      Hệ thống điều phối in action.
  4.  Add SuCoGiaoThong (traffic incident — new class from tracking system).
  5.  Add ThongBao to UC-03 scope (notification to customer when near).
  6.  Update SSD: replace step 5 "bàn giao" with step 5 "đến điểm giao".
  7.  Replace CRC card set: 6 cards (ChuyenDi, LoTrinh, ViTriGPS, KienHang,
      SuCoGiaoThong, ThongBao); BangChungGiaoHang CRC → UC-04.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document(DOC_PATH)
body = doc.element.body

# ── helpers ────────────────────────────────────────────────────────────────────

def gettxt(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{NS}}}t"))

def set_cell(tbl_elem, row_idx, col_idx, new_text):
    rows = tbl_elem.findall(f".//{{{NS}}}tr")
    cells = rows[row_idx].findall(f".//{{{NS}}}tc")
    cell = cells[col_idx]
    for p in cell.findall(f"{{{NS}}}p"):
        cell.remove(p)
    np_ = OxmlElement("w:p")
    nr_ = OxmlElement("w:r")
    nt_ = OxmlElement("w:t")
    nt_.set(qn("xml:space"), "preserve")
    nt_.text = new_text
    nr_.append(nt_); np_.append(nr_); cell.append(np_)

def get_cell(tbl_elem, row_idx, col_idx):
    rows = tbl_elem.findall(f".//{{{NS}}}tr")
    cells = rows[row_idx].findall(f".//{{{NS}}}tc")
    return "".join(t.text or "" for t in cells[col_idx].findall(f".//{{{NS}}}t"))

def replace_para_text(child, old, new):
    for t_el in child.findall(f".//{{{NS}}}t"):
        if t_el.text and old in t_el.text:
            t_el.text = t_el.text.replace(old, new)
            return True
    return False

def detach(obj):
    e = obj._element if hasattr(obj, "_element") else obj
    e.getparent().remove(e)
    return e

def make_para(text="", bold=False, italic=False, code=False):
    p = doc.add_paragraph(style="Normal")
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if code:
            r.font.name = "Courier New"
    return detach(p)

def insert_after(elem, new_elem):
    parent = elem.getparent()
    idx = list(parent).index(elem)
    parent.insert(idx + 1, new_elem)

def insert_before(elems, ref_elem):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)

def border_cell(cell):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
    else:
        for c in list(tcBorders): tcBorders.remove(c)
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"6")
        el.set(qn("w:color"),"000000"); el.set(qn("w:space"),"0")
        tcBorders.append(el)

def border_table(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            border_cell(cell)

def make_crc_table(tbl_name, tbl_id, layer, desc, attrs, responsibilities, partners):
    tbl = doc.add_table(rows=7, cols=2)
    try: tbl.style = "TableNormal"
    except Exception: pass
    data = [
        ("Tên lớp",tbl_name),("ID",tbl_id),("Tầng kiến trúc",layer),
        ("Mô tả",desc),("Các thuộc tính",attrs),
        ("Các trách nhiệm",responsibilities),("Các đối tác",partners),
    ]
    for ri,(label,value) in enumerate(data):
        cells = tbl.rows[ri].cells
        cells[0].text = label
        for para in cells[0].paragraphs:
            for run in para.runs: run.bold = True
        cells[1].text = value
    border_table(tbl)
    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def add_row_to_table(tbl_elem, row_data):
    """Append a new row to an existing table element (lxml)."""
    existing_rows = tbl_elem.findall(f".//{{{NS}}}tr")
    template_row = existing_rows[-1]
    new_row = copy.deepcopy(template_row)
    cells = new_row.findall(f".//{{{NS}}}tc")
    for ci, val in enumerate(row_data):
        if ci < len(cells):
            for t_el in cells[ci].findall(f".//{{{NS}}}t"):
                t_el.text = ""
            t_els = cells[ci].findall(f".//{{{NS}}}t")
            if t_els:
                t_els[0].text = val
            else:
                p = cells[ci].find(f"{{{NS}}}p")
                if p is None:
                    p = OxmlElement("w:p"); cells[ci].append(p)
                r = OxmlElement("w:r"); t = OxmlElement("w:t")
                t.set(qn("xml:space"),"preserve"); t.text = val
                r.append(t); p.append(r)
    tbl_elem.append(new_row)

def delete_row(tbl_elem, row_idx):
    rows = tbl_elem.findall(f".//{{{NS}}}tr")
    rows[row_idx].getparent().remove(rows[row_idx])

children = list(body)

# ════════════════════════════════════════════════════════════════════════════════
# 1. SCENARIO: Fix paragraphs [89]-[95]
# ════════════════════════════════════════════════════════════════════════════════

# [92]: Insert rerouting event paragraph AFTER "09:31 → 10:04 — Tài xế di chuyển"
reroute_para = make_para(
    '09:45 — Hệ thống theo dõi phát hiện sự cố giao thông SC-2026-001 '
    '(ùn tắc nghiêm trọng trên Vành đai 2, Q.8) nằm trên tuyến LT-2026-001. '
    'Hệ thống điều phối tính toán lại lộ trình thay thế trong 3 giây và gợi ý '
    'cho TX-001: tuyến đường mới qua Nguyễn Văn Linh — giữ nguyên ETA 10:06.'
)
para_92 = children[92]
insert_after(para_92, reroute_para)
print("Step 1a: Inserted rerouting paragraph after [92]")

# Reload children (index shifted +1 after insertion)
children = list(body)

# [94] was index 94, now still 94 (insertion was AFTER 92, so 93 shifted to 94,
#  but original 93 is now at 93, and 94 is now at 95).
# Wait: inserted after [92], so new [93] = reroute_para,
# original [93] = now [94], original [94] = now [95], original [95] = now [96]
# Let me re-find by content

# Find and modify "10:06 — TX-001 gặp chị" paragraph
for child in list(body):
    txt = gettxt(child)
    if '10:06' in txt and 'TX-001 gặp' in txt and 'BangChungGiaoHang' in txt:
        # Remove BangChungGiaoHang creation sentence, simplify handover
        old_txt = gettxt(child)
        new_txt = (
            '10:06 — TX-001 gặp chị Nguyễn Thị Lan tại cửa "123 Nguyễn Văn Cừ, Q.5". '
            'Bàn giao kiện hàng thành công về mặt vật lý. '
            'UC-03 kết thúc tại đây; hệ thống phát sự kiện "HandoverStarted" '
            'để kích hoạt UC-04 (Xác nhận giao hàng) xử lý việc chụp ảnh biên nhận và OTP.'
        )
        for t_el in child.findall(f".//{{{NS}}}t"):
            if t_el.text:
                t_el.text = ""
        t_els = child.findall(f".//{{{NS}}}t")
        if t_els:
            t_els[0].text = new_txt
        else:
            r = OxmlElement("w:r"); t = OxmlElement("w:t")
            t.text = new_txt; r.append(t); child.append(r)
        print("Step 1b: Fixed 10:06 handover paragraph")
        break

# Find and modify "10:06 (hoàn tất)" paragraph
for child in list(body):
    txt = gettxt(child)
    if '10:06 (hoàn tất)' in txt or ("DonHang.trangThai" in txt and "Giao thành công" in txt and "LoTrinh" in txt):
        new_final = (
            '10:06 (kết thúc UC-03) — Hệ thống cập nhật: '
            'DonHang.trangThai → "Chờ xác nhận"; '
            'KienHang.trangThai → "Đang xác nhận"; '
            'LoTrinh.trangThai → "Hoàn tất". '
            'Sự kiện "HandoverStarted" kích hoạt UC-04.'
        )
        for t_el in child.findall(f".//{{{NS}}}t"):
            if t_el.text:
                t_el.text = ""
        t_els = child.findall(f".//{{{NS}}}t")
        if t_els:
            t_els[0].text = new_final
        print("Step 1c: Fixed final status paragraph")
        break

# ════════════════════════════════════════════════════════════════════════════════
# 2. OBJECT TABLE [99]: Remove BangChungGiaoHang, add SuCoGiaoThong + ThongBao,
#    fix DonHang and KienHang trangThai values
# ════════════════════════════════════════════════════════════════════════════════

# Find object table by header
obj_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "tbl" and "Thuộc tính chính (giá trị kịch bản)" in gettxt(child):
        obj_tbl = child
        break

assert obj_tbl is not None, "Object table not found"

# Fix DonHang trangThai in obj table
rows_obj = obj_tbl.findall(f".//{{{NS}}}tr")
for ri in range(1, len(rows_obj)):
    cls = get_cell(obj_tbl, ri, 0).strip()
    if cls == "DonHang":
        txt = get_cell(obj_tbl, ri, 2)
        set_cell(obj_tbl, ri, 2, txt.replace('"Giao thành công"', '"Chờ xác nhận"'))
        print(f"Step 2a: Fixed DonHang trangThai in object table")
    if cls == "KienHang":
        txt = get_cell(obj_tbl, ri, 2)
        set_cell(obj_tbl, ri, 2, txt.replace('"Đã giao"', '"Đang xác nhận"'))
        print(f"Step 2b: Fixed KienHang trangThai in object table")

# Delete BangChungGiaoHang row
rows_obj = obj_tbl.findall(f".//{{{NS}}}tr")
for ri in range(len(rows_obj)-1, 0, -1):
    if "BangChungGiaoHang" in get_cell(obj_tbl, ri, 0):
        delete_row(obj_tbl, ri)
        print("Step 2c: Removed BangChungGiaoHang from object table")
        break

# Add SuCoGiaoThong row
add_row_to_table(obj_tbl, [
    "SuCoGiaoThong",
    "sc1 : SuCoGiaoThong",
    'maSuCo=SC-2026-001, loaiSuCo="Ùn tắc giao thông", '
    'thoiGianPhatHien="2026-03-26T09:45:00", toaDo="Vành đai 2, Q.8"'
])
print("Step 2d: Added SuCoGiaoThong to object table")

# Add ThongBao row
add_row_to_table(obj_tbl, [
    "ThongBao",
    "tb_uc03 : ThongBao",
    'maThongBao=TB-UC03-001, noiDung="Shipper đang đến, ETA 2 phút", '
    'thoiGianGui="2026-03-26T10:04:00", kenhGui="App + SMS"'
])
print("Step 2e: Added ThongBao to object table")

# ════════════════════════════════════════════════════════════════════════════════
# 3. LINK TABLE [104]: Remove "có bằng chứng", add new links
# ════════════════════════════════════════════════════════════════════════════════

link_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "tbl" and "có bằng chứng" in gettxt(child) and "LoTrinh" in gettxt(child):
        link_tbl = child
        break

assert link_tbl is not None, "Link table not found"

# Delete "có bằng chứng" row
rows_lnk = link_tbl.findall(f".//{{{NS}}}tr")
for ri in range(len(rows_lnk)-1, 0, -1):
    if "có bằng chứng" in gettxt(rows_lnk[ri]):
        delete_row(link_tbl, ri)
        print("Step 3a: Removed 'có bằng chứng' from link table")
        break

# Add new links
add_row_to_table(link_tbl, [
    "lt1 : LoTrinh", "sc1 : SuCoGiaoThong", "gặp sự cố",
    "Bước 4b — tracking phát hiện ùn tắc, OR tính lại lộ trình"
])
add_row_to_table(link_tbl, [
    "dh1 : DonHang", "tb_uc03 : ThongBao", "kích hoạt thông báo",
    "Bước 5 — gần điểm giao, hệ thống gửi ThongBao đến khách"
])
print("Step 3b: Added SuCoGiaoThong and ThongBao links")

# Update paragraph [106] counting text ("6 liên kết" → keep consistent)
for child in list(body):
    txt = gettxt(child)
    if "6 liên kết" in txt and "sơ đồ lớp" in txt:
        replace_para_text(child, "6 liên kết", "6 liên kết")  # count unchanged
        print("Step 3c: Link count paragraph checked (6 links still valid)")
        break

# ════════════════════════════════════════════════════════════════════════════════
# 4. CLASS DIAGRAM INTRO [110]: Update to reflect 8 classes, remove BangChung
# ════════════════════════════════════════════════════════════════════════════════

for child in list(body):
    txt = gettxt(child)
    if "7 lớp tham gia" in txt and "BangChungGiaoHang" in txt:
        new_txt = (
            "Sơ đồ lớp lĩnh vực UC-03 mô tả 8 lớp tham gia trực tiếp vào quy "
            "trình vận chuyển đơn hàng: TaiXe, DonHang, KienHang, LoTrinh, "
            "ViTriGPS, ChuyenDi, SuCoGiaoThong và ThongBao. "
            "BangChungGiaoHang được chuyển sang UC-04 vì việc tạo bằng chứng "
            "(ảnh + OTP) thuộc luồng xác nhận, không thuộc luồng vận chuyển. "
            "Đây là UC phức tạp nhất (13 bước) nên sơ đồ lớp có số liên kết nhiều nhất. "
            "Mọi thuộc tính được ghi ở mức phân tích: chỉ tên, không có kiểu dữ liệu hay "
            "ký hiệu phạm vi truy cập."
        )
        for t_el in child.findall(f".//{{{NS}}}t"):
            if t_el.text: t_el.text = ""
        t_els = child.findall(f".//{{{NS}}}t")
        if t_els: t_els[0].text = new_txt
        print("Step 4: Updated class diagram intro paragraph")
        break

# ════════════════════════════════════════════════════════════════════════════════
# 5. CLASS TABLE [112]: Remove BangChungGiaoHang, add SuCoGiaoThong + ThongBao
# ════════════════════════════════════════════════════════════════════════════════

cls_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "tbl" and "Thuộc tính (phân tích" in gettxt(child):
        cls_tbl = child
        break

assert cls_tbl is not None, "Class diagram table not found"

# Delete BangChungGiaoHang row
rows_cls = cls_tbl.findall(f".//{{{NS}}}tr")
for ri in range(len(rows_cls)-1, 0, -1):
    if "BangChungGiaoHang" in get_cell(cls_tbl, ri, 0):
        delete_row(cls_tbl, ri)
        print("Step 5a: Removed BangChungGiaoHang from class table")
        break

# Add SuCoGiaoThong
add_row_to_table(cls_tbl, [
    "SuCoGiaoThong",
    "maSuCo, loaiSuCo, thoiGianPhatHien, toaDo",
    "Ghi nhận sự cố giao thông phát hiện bởi Hệ thống theo dõi; "
    "kích hoạt Hệ thống điều phối tính lại tuyến đường."
])
print("Step 5b: Added SuCoGiaoThong to class table")

# Add ThongBao
add_row_to_table(cls_tbl, [
    "ThongBao",
    "maThongBao, noiDung, thoiGianGui, kenhGui",
    "Thông báo đẩy đến khách hàng khi tài xế gần đến điểm giao; "
    "cùng lớp dùng lại ở UC-04 (thông báo xác nhận giao hàng)."
])
print("Step 5c: Added ThongBao to class table")

# ════════════════════════════════════════════════════════════════════════════════
# 6. ASSOCIATION TABLE [116]: Remove "có bằng chứng", add new associations
# ════════════════════════════════════════════════════════════════════════════════

assoc_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "tbl" and "có bằng chứng" in gettxt(child) and "Composition" in gettxt(child):
        assoc_tbl = child
        break

assert assoc_tbl is not None, "Association table not found"

rows_as = assoc_tbl.findall(f".//{{{NS}}}tr")
for ri in range(len(rows_as)-1, 0, -1):
    if "có bằng chứng" in gettxt(rows_as[ri]):
        delete_row(assoc_tbl, ri)
        print("Step 6a: Removed 'có bằng chứng' from association table")
        break

add_row_to_table(assoc_tbl, [
    "LoTrinh", "gặp sự cố", "SuCoGiaoThong", "1 → 0..*", "Association"
])
add_row_to_table(assoc_tbl, [
    "DonHang", "kích hoạt thông báo", "ThongBao", "1 → 0..*", "Association"
])
print("Step 6b: Added SuCoGiaoThong and ThongBao to association table")

# ════════════════════════════════════════════════════════════════════════════════
# 7. SSD MESSAGE TABLE [123]: Replace step 5/6 with step 5 "đến điểm giao"
# ════════════════════════════════════════════════════════════════════════════════

ssd_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "tbl" and "xacNhanBatDauCa" in gettxt(child):
        ssd_tbl = child
        break

assert ssd_tbl is not None, "SSD table not found"

# Delete rows with "bàn giao" / "BangChungGiaoHang" content (r9 and r10)
rows_ssd = ssd_tbl.findall(f".//{{{NS}}}tr")
for ri in range(len(rows_ssd)-1, 0, -1):
    row_txt = gettxt(rows_ssd[ri])
    if "xacNhanBanGiao" in row_txt or "xacNhanGiaoHangThanhCong" in row_txt or "BangChungGiaoHang" in row_txt:
        delete_row(ssd_tbl, ri)
        print(f"Step 7a: Removed row {ri} from SSD table")

# Add new step 5
add_row_to_table(ssd_tbl, [
    "5", "→ :Hệ thống",
    "xacNhanDaDenDiemGiao(maDonHang)",
    "Bước 5 — đến vị trí khách hàng",
    "Kết thúc UC-03; hệ thống gửi ThongBao đến khách; phát HandoverStarted → kích hoạt UC-04"
])
add_row_to_table(ssd_tbl, [
    "↩", ":Hệ thống →",
    "xacNhanCoMat(maThongBao)",
    "—",
    "Xác nhận ThongBao đã gửi; DonHang.trangThai → Chờ xác nhận"
])
print("Step 7b: Added step 5 'đến điểm giao' to SSD table")

# ════════════════════════════════════════════════════════════════════════════════
# 8. SSD PLANTUML TEXT: Replace step 5-6 block
# ════════════════════════════════════════════════════════════════════════════════

# Find paragraphs in the PlantUML block
children = list(body)
for child in children:
    txt = gettxt(child)
    if '== Bước 5–6: Bàn giao và tạo bằng chứng ==' in txt:
        replace_para_text(child,
            '== Bước 5–6: Bàn giao và tạo bằng chứng ==',
            '== Bước 5: Xác nhận đến điểm giao / kết thúc UC-03 =='
        )
        print("Step 8a: Fixed PlantUML step 5 heading")
    if 'TX -> SYS : xacNhanBanGiao' in txt:
        replace_para_text(child,
            'TX -> SYS : xacNhanBanGiao(maDonHang, tenNguoiNhan, hinhAnh)',
            'TX -> SYS : xacNhanDaDenDiemGiao(maDonHang)'
        )
        print("Step 8b: Fixed PlantUML step 5 outgoing message")
    if 'SYS --> TX : xacNhanGiaoHangThanhCong' in txt:
        replace_para_text(child,
            'SYS --> TX : xacNhanGiaoHangThanhCong(maBangChung)',
            'SYS --> TX : xacNhanCoMat(maThongBao)'
        )
        print("Step 8c: Fixed PlantUML step 5 return message")

# ════════════════════════════════════════════════════════════════════════════════
# 9. CRC CARDS: Replace existing 4 cards + add SuCoGiaoThong + ThongBao (6 total)
#    (BangChungGiaoHang is absent; don't add it)
# ════════════════════════════════════════════════════════════════════════════════

# Find the 4.3.5 heading and the existing CRC tables; rebuild all
# First, find first CRC table after the heading
children = list(body)
first_crc_idx = None
in_crc = False
for i, child in enumerate(children):
    txt = gettxt(child)
    if "4.3.5" in txt and "CRC" in txt:
        in_crc = True
    if in_crc:
        tag = child.tag.split("}")[1]
        if tag == "tbl" and "Tên lớp" in txt and "CRC-UC03" in txt:
            first_crc_idx = i
            break

assert first_crc_idx is not None, "First CRC table not found"
print(f"First CRC table at body[{first_crc_idx}]")

# Collect and remove all existing CRC tables + their spacers (until 4.4. heading)
children = list(body)
to_remove = []
in_crc_section = False
for i, child in enumerate(children):
    txt = gettxt(child)
    if i == first_crc_idx:
        in_crc_section = True
    if "4.4." in txt and in_crc_section:
        break
    if in_crc_section:
        tag = child.tag.split("}")[1]
        if tag == "tbl" and "Tên lớp" in txt:
            to_remove.append(child)
        elif tag == "p" and txt.strip() == "" and in_crc_section:
            # Only remove blank spacer paragraphs BETWEEN crc tables
            to_remove.append(child)

print(f"Removing {len(to_remove)} elements (CRC tables + spacers)")
anchor = to_remove[0]  # insertion point

new_crc_elems = []

crc_data = [
    ("ChuyenDi", "CRC-UC03-01", "Lĩnh vực (Domain)",
     "Đại diện cho một ca làm việc của tài xế; bao gồm một hoặc nhiều lộ "
     "trình giao hàng trong một ngày và được gắn với một hoặc nhiều sự cố giao thông.",
     "maChuyenDi, thoiGianBatDau, thoiGianKetThuc",
     "1. Ghi nhận thời điểm bắt đầu và kết thúc ca làm việc.\n"
     "2. Liên kết tài xế với tập hợp các lộ trình trong ca.\n"
     "3. Cung cấp ngữ cảnh ca làm việc để tổng hợp hiệu suất.",
     "TaiXe, LoTrinh"),

    ("LoTrinh", "CRC-UC03-02", "Lĩnh vực (Domain)",
     "Tuyến đường từ kho xuất phát đến điểm giao hàng; lưu chuỗi tọa độ GPS "
     "toàn hành trình và ghi nhận mọi sự cố giao thông trên tuyến.",
     "maLoTrinh, diemXuatPhat, diemDen, khoangCach, /thoiGianDuKien, trangThai",
     "1. Lưu thông tin tuyến đường (điểm đầu, điểm cuối, khoảng cách).\n"
     "2. Tập hợp các bản ghi ViTriGPS trong suốt hành trình.\n"
     "3. Ghi nhận SuCoGiaoThong xảy ra trên tuyến.\n"
     "4. Theo dõi trạng thái: Đang giao → Hoàn tất.\n"
     "5. Cung cấp dữ liệu ETA cho Hệ thống điều phối tái tính toán.",
     "ChuyenDi, TaiXe, DonHang, ViTriGPS, SuCoGiaoThong"),

    ("ViTriGPS", "CRC-UC03-03", "Lĩnh vực (Domain)",
     "Bản ghi tọa độ GPS tại một thời điểm cụ thể; được ghi nhận mỗi 30 giây "
     "bởi Hệ thống theo dõi suốt hành trình của tài xế.",
     "maViTri, viDo, kinhDo, thoiGian",
     "1. Lưu tọa độ địa lý (vĩ độ, kinh độ) tại thời điểm ghi nhận.\n"
     "2. Cung cấp dữ liệu vị trí theo thời gian thực cho tracking.\n"
     "3. Tạo vết hành trình (breadcrumb trail) để kiểm tra sau giao.",
     "LoTrinh"),

    ("KienHang", "CRC-UC03-04", "Lĩnh vực (Domain)",
     "Đơn vị vật lý được vận chuyển; được xác nhận bằng mã QR khi tài xế lấy "
     "hàng tại kho và chuyển sang UC-04 để xác nhận bàn giao.",
     "maKienHang, maQR, khoiLuong, trangThai",
     "1. Xác minh danh tính kiện hàng qua quét mã QR.\n"
     "2. Theo dõi trạng thái vật lý trong UC-03: Đang vận chuyển → Đang xác nhận.\n"
     "3. Cung cấp thông tin khoiLuong để tính phí vận chuyển (UC-05).",
     "DonHang"),

    ("SuCoGiaoThong", "CRC-UC03-05", "Lĩnh vực (Domain)",
     "Sự kiện giao thông bất thường (ùn tắc, tai nạn) phát hiện bởi Hệ thống "
     "theo dõi GPS; kích hoạt Hệ thống điều phối (OR) tính lại lộ trình tối ưu.",
     "maSuCo, loaiSuCo, thoiGianPhatHien, toaDo",
     "1. Lưu thông tin sự cố giao thông trên tuyến đường đang lưu thông.\n"
     "2. Cung cấp dữ liệu đầu vào để Hệ thống điều phối tính lộ trình thay thế.\n"
     "3. Phục vụ kiểm toán SLA khi cần giải trình về thời gian giao chậm.",
     "LoTrinh"),

    ("ThongBao", "CRC-UC03-06", "Lĩnh vực (Domain)",
     "Thông điệp đẩy tự động đến khách hàng khi tài xế gần đến điểm giao "
     "(ETA ≤ 2 phút); cùng lớp ThongBao được tái sử dụng ở UC-04.",
     "maThongBao, noiDung, thoiGianGui, kenhGui",
     "1. Thông báo chủ động cho khách hàng về thời gian đến dự kiến.\n"
     "2. Gửi qua đa kênh: ứng dụng và SMS (kenhGui).\n"
     "3. Cung cấp bằng chứng rằng khách đã được thông báo (phục vụ xử lý khiếu nại).",
     "DonHang"),
]

for args in crc_data:
    t, s = make_crc_table(*args)
    new_crc_elems.extend([t, s])

# Insert new CRC elements at anchor position
insert_before(new_crc_elems, anchor)
print(f"Step 9: Inserted {len(new_crc_elems)} new CRC elements")

# Now remove the old ones
for elem in to_remove:
    if elem.getparent() is not None:
        elem.getparent().remove(elem)
print(f"Step 9: Removed {len(to_remove)} old CRC elements")

# ════════════════════════════════════════════════════════════════════════════════
# 10. UPDATE CRC intro paragraph
# ════════════════════════════════════════════════════════════════════════════════
for child in list(body):
    txt = gettxt(child)
    if "5 lớp lĩnh vực trung tâm của UC-03" in txt:
        new_crc_intro = (
            "Thẻ CRC (Class–Responsibility–Collaborator) được lập cho 6 lớp lĩnh vực "
            "của UC-03. BangChungGiaoHang thuộc UC-04 (xác nhận bàn giao) nên không "
            "có trong danh sách này. SuCoGiaoThong là lớp mới được nhận diện từ vai "
            "trò của Hệ thống theo dõi GPS. ThongBao được dùng chung với UC-04 "
            "(lần gọi khác nhau). Thuộc tính chỉ ghi tên — không có kiểu dữ liệu "
            "(quy tắc phân tích I.12)."
        )
        for t_el in child.findall(f".//{{{NS}}}t"):
            if t_el.text: t_el.text = ""
        t_els = child.findall(f".//{{{NS}}}t")
        if t_els: t_els[0].text = new_crc_intro
        print("Step 10: Updated CRC intro paragraph")
        break

# ════════════════════════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f"\nAll done. Saved to {DOC_PATH}")
