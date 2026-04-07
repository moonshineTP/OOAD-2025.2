"""
Creates Project/Ch5_MoHinhHoaHanhVi.docx
Chapter V: Mô hình hóa hành vi — LogiFast delivery system
Based on docs/week30-plan.md structure (Section 1-6)

Run from repo root:
    .venv\Scripts\python.exe scripts\create_ch5_full.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("Project/Ch4_MoHinhHoaCauTruc.docx")   # style donor
OUT = Path("Project/Ch5_MoHinhHoaHanhVi.docx")

# ══════════════════════════════════════════════════════════════════════════════
# BOOTSTRAP: copy Ch4 for styles, then clear body
# ══════════════════════════════════════════════════════════════════════════════
shutil.copy(SRC, OUT)
doc = Document(str(OUT))

body = doc.element.body
children = list(body)
sect_pr = body.find(qn("w:sectPr"))
for child in children:
    if child is not sect_pr:
        body.remove(child)

# ══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def set_borders(table):
    """Add single black borders to table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def shade_cells(row, hex_color="D9D9D9"):
    """Shade row with gray background."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for existing in tcPr.findall(qn("w:shd")):
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def set_cell_font(cell, size_pt, bold=False):
    """Set font size and bold for all runs in cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size_pt)
            if bold:
                run.bold = True

def set_cell_width(cell, twips):
    """Set cell width in twips (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def make_table(headers, rows, font_size=10, header_bold=True):
    """Create table with borders, gray header, fixed font size."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_borders(tbl)
    # Header
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        set_cell_font(tbl.rows[0].cells[i], font_size, bold=header_bold)
    shade_cells(tbl.rows[0])
    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            tbl.rows[ri + 1].cells[ci].text = str(val) if val else ""
            set_cell_font(tbl.rows[ri + 1].cells[ci], font_size)
    return tbl

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def add_placeholder(text):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    return p

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph("V. Mô hình hóa hành vi", style="Title")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — GIỚI THIỆU
# ══════════════════════════════════════════════════════════════════════════════

add_h1("1. Giới thiệu")

add_para(
    "Sau khi mô hình cấu trúc (Chương IV) xác định những gì hệ thống ghi nhớ — "
    "các lớp lĩnh vực, thuộc tính và mối quan hệ tĩnh giữa chúng — mô hình hóa "
    "hành vi trả lời câu hỏi như thế nào các đối tượng đó biến đổi trạng thái và "
    "tương tác ra sao với nhau trong từng ca sử dụng. Theo bài giảng Ch04 "
    "(TS. Nguyễn Bá Ngọc), ba vấn đề cốt lõi cần mô hình hóa là: (1) sự biến đổi "
    "trạng thái của đối tượng theo tiến trình nghiệp vụ; (2) tương tác giữa tác "
    "nhân và hệ thống; (3) tương tác giữa các đối tượng nội bộ để đáp ứng hoạt "
    "động nghiệp vụ."
)

add_para(
    "Chương V trình bày bốn nội dung theo đúng cấu trúc bài giảng: sơ đồ máy "
    "trạng thái cho các đối tượng trọng tâm (Mục 2); sơ đồ tuần tự mức hệ thống "
    "— SSD — cho cả 5 ca sử dụng (Mục 3); sơ đồ tuần tự mức nghiệp vụ và sơ đồ "
    "giao tiếp mức nghiệp vụ cho UC-03 (Mục 4); và ma trận CRUD(E) tổng hợp "
    "(Mục 5). Các sơ đồ này là cơ sở trực tiếp cho thiết kế chi tiết ở Phần II."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SƠ ĐỒ MÁY TRẠNG THÁI
# ══════════════════════════════════════════════════════════════════════════════

add_h1("2. Sơ đồ máy trạng thái")

add_para(
    "Theo tiêu chí lựa chọn từ bài giảng (slide 5), đối tượng được đưa vào mô "
    "hình máy trạng thái phải thỏa đồng thời: (a) xuất hiện trong nhiều ca sử "
    "dụng, (b) có trạng thái liên quan đến công việc nghiệp vụ được biểu diễn "
    "qua thuộc tính rõ ràng, và (c) có sự thay đổi trạng thái được kích hoạt bởi "
    "các sự kiện nghiệp vụ xác định. Trong LogiFast, ba đối tượng thỏa điều kiện: "
    "DonHang (xuyên suốt 5 UC với thuộc tính trangThai), TaiXe (điều kiện bảo vệ "
    "phân công trong UC-02), và LoTrinh (tái tối ưu động bởi OR engine trong "
    "UC-03 — trạng thái duy nhất có do-activity liên tục trong toàn hệ thống)."
)

# Terminology table
make_table(
    headers=["Thuật ngữ", "Mô tả"],
    rows=[
        ["Trạng thái", "Một giai đoạn trong tiến trình; bộ giá trị thuộc tính thỏa điều kiện nhất định"],
        ["Bước chuyển (Transition)", "Từ trạng thái nguồn → trạng thái đích; nhãn: sự_kiện [điều_kiện] / hành_động"],
        ["Sự kiện kích hoạt", "Nguyên nhân dẫn đến chuyển trạng thái"],
        ["Điều kiện bảo vệ", "Chỉ chuyển nếu điều kiện [...] được đáp ứng"],
        ["Biểu thức hành vi", "Được thực hiện và hoàn thành trước khi chuyển sang trạng thái đích"],
        ["Hoạt động bên trong (do-activity)", "Chạy liên tục trong khi đối tượng ở trạng thái đó; hoàn thành khi kết thúc"],
        ["Trạng thái tổng hợp", "Chứa máy trạng thái con với trạng thái bắt đầu và kết thúc riêng"],
    ],
    font_size=10,
)
spacer()

# ──────────────────────────────────────────────────────────────────────────────
# 2.1 TaiXe
# ──────────────────────────────────────────────────────────────────────────────

add_h2("2.1. Sơ đồ máy trạng thái — TaiXe")

add_para(
    "TaiXe có vòng đời theo ca làm việc. Trạng thái của tài xế là điều kiện bảo "
    "vệ quan trọng trong UC-02: hệ thống phân công chỉ xét các tài xế đang ở "
    "trạng thái \"Sẵn sàng nhận đơn\". Ngoài ra, trạng thái \"Đang giao hàng\" "
    "giải thích tại sao cùng một tài xế không thể nhận thêm đơn mới trong UC-02 "
    "khi đang thực hiện UC-03. Vòng lặp Sẵn sàng nhận đơn ↔ Đang giao hàng có "
    "thể lặp lại nhiều lần trong một ca, tương ứng với số đơn hàng được giao."
)

make_table(
    headers=["Từ trạng thái", "Sự kiện", "Điều kiện", "Hành động", "Đến trạng thái", "UC / Bước", "Thông điệp SSD"],
    rows=[
        ["●", "—", "—", "—", "Ngoài ca", "—", "—"],
        ["Ngoài ca", "ShiftStarted", "—", "tạo ChuyenDi", "Sẵn sàng nhận đơn", "UC-03 B1", "xacNhanBatDauCa(maTaiXe)"],
        ["Sẵn sàng nhận đơn", "AssignmentReceived", "—", "cập nhật trangThai", "Đang giao hàng", "UC-02 B3", "phanCongTaiXe(...)"],
        ["Đang giao hàng", "DeliveryCompleted", "—", "cập nhật trangThai", "Sẵn sàng nhận đơn", "UC-03 B6", "xacNhanGiaoHangThanhCong(...)"],
        ["Sẵn sàng nhận đơn", "ShiftEnded", "—", "ghi thoiGianKetThuc", "Kết thúc ca", "UC-03 (end)", "ketThucCa(maTaiXe)"],
        ["Kết thúc ca", "—", "—", "—", "⊙", "—", "—"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ máy trạng thái TaiXe — Project/Diagrams/TaiXe_StateMachine.puml")

# ──────────────────────────────────────────────────────────────────────────────
# 2.2 LoTrinh
# ──────────────────────────────────────────────────────────────────────────────

add_h2("2.2. Sơ đồ máy trạng thái — LoTrinh")

add_para(
    "LoTrinh là đối tượng kỹ thuật phức tạp nhất trong UC-03. Không giống các "
    "lớp khác, LoTrinh không có vòng đời tuyến tính: nó có thể bị tái tính toán "
    "(recalculated) bởi OR engine khi phát sinh sự kiện bất thường — tắc đường "
    "được phát hiện bởi Hệ thống theo dõi, hoặc GPS tài xế lệch khỏi tuyến đã "
    "tính hơn 500m. Trạng thái \"Đang tái tối ưu\" là trạng thái duy nhất trong "
    "toàn bộ mô hình LogiFast có hoạt động nội bộ liên tục (do / OR engine "
    "recalculates route), thể hiện tính chủ động của hệ thống."
)

add_para(
    "Điều kiện bảo vệ [GPS lệch >500m] được cấp trực tiếp từ thông điệp "
    "\"Cập nhật vị trí GPS định kỳ()\" trong vòng lặp của SSD UC-03 (Mục 3.3). "
    "Đây là liên kết I.18: chuyển trạng thái Đang thực hiện → Đang tái tối ưu "
    "có thể truy xuất về cả bước UC-03 Bước 4 lẫn thông điệp SSD. Thuộc tính "
    "/thoiGianDuKien được tái tính mỗi lần OR hoàn thành, chứng minh "
    "thoiGianDuKien là thuộc tính suy diễn động (derived), không tĩnh — ký hiệu "
    "/ theo quy ước UML."
)

make_table(
    headers=["Từ trạng thái", "Sự kiện", "Điều kiện", "Hành động", "Đến trạng thái", "UC / Bước", "Thông điệp SSD"],
    rows=[
        ["●", "—", "—", "OR engine tính tuyến", "Đã tính toán", "UC-03 B1", "Bắt đầu ca làm việc()"],
        ["Đã tính toán", "DriverDeparted", "[tài xế xác nhận]", "ghi thoiGianBatDau", "Đang thực hiện", "UC-03 B3", "Hiển thị lộ trình và ETA()"],
        ["Đang thực hiện", "RouteDeviated", "[GPS lệch >500m OR tắc đường]", "kích hoạt OR engine", "Đang tái tối ưu", "UC-03 B4", "Cập nhật vị trí GPS định kỳ()"],
        ["Đang tái tối ưu", "OptimizationDone", "[tuyến mới tốt hơn ≥10%]", "cập nhật khoangCach, /ETA", "Đang thực hiện", "UC-03 B4 loop", "Trả kết quả lộ trình mới()"],
        ["Đang tái tối ưu", "OptimizationDone", "[tuyến cũ vẫn tối ưu]", "giữ nguyên", "Đang thực hiện", "UC-03 B4 loop", "—"],
        ["Đang thực hiện", "PackageHandedOver", "[GPS ±200m điểm đến]", "ghi thoiGianKetThuc", "Hoàn tất", "UC-03 B6", "Xác nhận bàn giao hàng()"],
        ["Đang thực hiện", "OrderCancelled", "—", "ghi lý do", "Bị hủy", "UC-03 alt", "—"],
        ["Hoàn tất", "—", "—", "—", "⊙", "—", "—"],
        ["Bị hủy", "—", "—", "—", "⊙", "—", "—"],
    ],
    font_size=8,
)
spacer()
add_placeholder("Chèn sơ đồ máy trạng thái LoTrinh — Project/Diagrams/LoTrinh_StateMachine.puml")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SƠ ĐỒ TUẦN TỰ MỨC HỆ THỐNG (SSD)
# ══════════════════════════════════════════════════════════════════════════════

add_h1("3. Sơ đồ tuần tự mức hệ thống (SSD)")

add_para(
    "Sơ đồ tuần tự mức hệ thống (System Sequence Diagram — SSD) có cấu trúc cố "
    "định: chỉ gồm tác nhân và hệ thống (\":Hệ thống LogiFast\"). Các thông điệp "
    "tương ứng với các hoạt động nghiệp vụ ở mức khái quát cao — mỗi thông điệp "
    "đại diện cho một bước trong đặc tả UC. SSD là đầu vào trực tiếp cho bước "
    "thiết kế: mỗi thông điệp gửi đến hệ thống sẽ trở thành một hợp đồng thông "
    "điệp (system operation contract) ở Phần II. Điều này đảm bảo tính nhất quán "
    "giữa mô hình phân tích và mô hình thiết kế."
)

add_para(
    "Quy tắc ký hiệu (slides 17–18, 23): tên đường sống viết dạng "
    "tênĐốiTượng:TênLớp, gạch chân, căn giữa. SSD chỉ có đúng hai hoặc nhiều "
    "đường sống: tác nhân (actor), các hệ thống ngoài (external systems) và "
    "\":Hệ thống LogiFast\". Mô tả thông điệp theo cú pháp: [biểuThứcLôGic] "
    "tênThôngĐiệp(). Điều kiện bảo vệ trong [...] xác định khi nào thông điệp "
    "được gửi. Khung kết hợp (combined fragments): loop (lặp), opt (tùy chọn), "
    "alt (rẽ nhánh), break (dừng vòng lặp)."
)

# ──────────────────────────────────────────────────────────────────────────────
# 3.1 UC-01 placeholder
# ──────────────────────────────────────────────────────────────────────────────

add_h2("3.1. SSD — UC-01: Đặt đơn hàng")
add_para("Tác nhân: Khách hàng. UC-01 mô tả luồng khách hàng xác nhận giỏ hàng, "
         "nhập địa chỉ giao, chọn hình thức thanh toán và đặt đơn.")
add_placeholder("Chèn SSD UC-01 — Project/Diagrams/UC01_SSD.puml  |  Owner: Trương Văn Hồng")
spacer()

# ──────────────────────────────────────────────────────────────────────────────
# 3.2 UC-02 placeholder
# ──────────────────────────────────────────────────────────────────────────────

add_h2("3.2. SSD — UC-02: Phân công giao hàng")
add_para("Tác nhân: Hệ thống phân công (tự động). UC-02 mô tả quy trình hệ thống "
         "tìm tài xế phù hợp, tạo phiếu phân công và thông báo tài xế chấp nhận.")
add_placeholder("Chèn SSD UC-02 — Project/Diagrams/UC02_SSD.puml  |  Owner: Nguyễn Quý Duy")
spacer()

# ──────────────────────────────────────────────────────────────────────────────
# 3.3 UC-03 — FULL (16 messages, Vietnamese parameterless, from week30-plan)
# ──────────────────────────────────────────────────────────────────────────────

add_h2("3.3. SSD — UC-03: Vận chuyển đơn hàng")

add_para(
    "Tác nhân: Tài xế giao hàng, Khách hàng, Hệ thống theo dõi, Hệ thống điều "
    "phối. UC-03 là ca sử dụng phức tạp nhất với 13 bước nghiệp vụ và sự tham "
    "gia của nhiều hệ thống ngoài. Tại mức hệ thống (SSD), mọi thông điệp phải "
    "là parameterless và là Tiếng Việt có dấu được đánh số thứ tự tuần tự. "
    "Các Actor hệ thống phụ trợ (Hệ thống theo dõi, Hệ thống điều phối) được "
    "biểu diễn thành Lifelines độc lập bên cạnh :Hệ thống LogiFast."
)

make_table(
    headers=["#", "Hướng", "Tên thông điệp (Tiếng Việt, Parameterless)", "Bước UC & Luồng", "Ghi chú"],
    rows=[
        ["1", "Tài xế → Hệ thống", "Bắt đầu ca làm việc()", "B1 (Luồng chính)", "Sinh sự kiện tạo lộ trình"],
        ["2", "Hệ thống → Tài xế", "Yêu cầu quét mã kiện hàng()", "B1 (Luồng chính)", "Hệ thống yêu cầu xác nhận kiện hàng"],
        ["3", "Tài xế → Hệ thống", "Quét mã kiện hàng()", "B1 (Luồng chính)", "Xác thực kiện hàng"],
        ["4", "Hệ thống → Hệ thống điều phối", "Tính toán lộ trình tối ưu()", "B1 (Luồng chính)", "Giao tiếp với External System phân luồng"],
        ["5", "Hệ thống điều phối → Hệ thống", "Trả kết quả lộ trình()", "B1 (Luồng chính)", "Trả về lộ trình tối ưu để hiển thị"],
        ["6", "Hệ thống → Tài xế", "Hiển thị lộ trình và ETA()", "B2 (Luồng chính)", "Tài xế nhận lộ trình"],
        ["7 (loop)", "Hệ thống theo dõi → Hệ thống", "Cập nhật vị trí GPS định kỳ()", "B2, B3 (Luồng chính)", "Ghi nhận Tracking mỗi 30 giây"],
        ["8 (alt)", "Hệ thống theo dõi → Hệ thống", "Thông báo phát hiện sự cố giao thông()", "B4 (Luồng thay thế d)", "Nếu phát hiện tắc đường"],
        ["9", "Hệ thống → Hệ thống điều phối", "Yêu cầu tính lại lộ trình()", "B4 (Luồng thay thế d)", "Yêu cầu tái tối ưu"],
        ["10", "Hệ thống điều phối → Hệ thống", "Trả kết quả lộ trình mới()", "B5 (Luồng thay thế d)", "Trả về lộ trình thay thế"],
        ["11", "Hệ thống → Tài xế", "Cập nhật lộ trình thay thế()", "B5 (Luồng thay thế d)", "Hiển thị lộ trình mới cho xe"],
        ["12", "Tài xế → Hệ thống", "Xác nhận đã đến điểm giao()", "B6 (Luồng chính)", "Cập nhật trạng thái 'Đã đến điểm giao'"],
        ["13", "Hệ thống → Khách hàng", "Gửi thông báo Tài xế đã đến()", "B6 (Luồng chính)", "Thông báo cho Khách hàng chuẩn bị nhận"],
        ["14", "Tài xế → Hệ thống", "Xác nhận bàn giao hàng()", "B7 (Luồng chính)", "Tài xế đẩy trạng thái giao xong"],
        ["15", "Khách hàng → Hệ thống", "Khách hàng đồng ý nhận hàng()", "B7 (Luồng chính)", "Hoàn tất flow vật lý"],
        ["16", "Hệ thống → Tài xế", "Thông báo hoàn tất đơn hàng()", "B7 (Luồng chính)", "Thông báo kết thúc tiến trình"],
    ],
    font_size=8,
)
spacer()
add_placeholder("Chèn SSD UC-03 — Project/Diagrams/UC03_SSD_system.drawio  |  Owner: Phạm Gia Hưng")
spacer()

# ──────────────────────────────────────────────────────────────────────────────
# 3.4 UC-04 placeholder
# ──────────────────────────────────────────────────────────────────────────────

add_h2("3.4. SSD — UC-04: Xác nhận giao hàng")
add_para("Tác nhân: Shipper (TaiXe). UC-04 mô tả luồng Shipper mở phiếu, yêu cầu "
         "OTP, nhận OTP từ khách hàng và xác nhận giao hàng thành công.")
add_placeholder("Chèn SSD UC-04 — Project/Diagrams/UC04_SSD.puml  |  Owner: Đinh Việt Hùng")
spacer()

# ──────────────────────────────────────────────────────────────────────────────
# 3.5 UC-05 placeholder
# ──────────────────────────────────────────────────────────────────────────────

add_h2("3.5. SSD — UC-05: Giao hàng hoàn tất / Thanh toán")
add_para("Tác nhân: Shipper (TaiXe). UC-05 mô tả luồng Shipper xem thông tin kết "
         "toán, xác nhận nhận tiền COD và thực hiện kết toán với hệ thống.")
add_placeholder("Chèn SSD UC-05 — Project/Diagrams/UC05_SSD.puml  |  Owner: Nguyễn Ngọc Toàn")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — SƠ ĐỒ TUẦN TỰ MỨC THIẾT KẾ / NGHIỆP VỤ
# ══════════════════════════════════════════════════════════════════════════════

add_h1("4. Sơ đồ tuần tự mức thiết kế / nghiệp vụ")

add_para(
    "Sơ đồ tuần tự mức thiết kế / nghiệp vụ (Business-Level / Real-Level "
    "Sequence Diagram) chuyển góc nhìn từ ranh giới hệ thống vào bên trong: "
    "thay vì chỉ hai đường sống (actor + :Hệ thống), sơ đồ này triển khai các "
    "đối tượng lĩnh vực cụ thể tham gia vào UC-03 — Boundary (:GiaoDienTaiXe), "
    "Controller (:VanChuyenController), và Entities (:ChuyenDi, :LoTrinh, "
    ":DonHang, :ViTriGPS, :SuCoGiaoThong, :BangChungGiaoHang, :KienHang). "
    "Mỗi mũi tên có thể truy xuất về một quan hệ trong sơ đồ lớp UC-03 (Ch IV "
    "§4.3.2) và một cộng tác viên trong thẻ CRC tương ứng (Ch IV §4.3.4)."
)

add_h2("4.1. Kịch bản thông điệp — UC-03 (Business-Level SD)")

add_para(
    "Các Lifeline tham gia: Tài xế, Khách hàng, Hệ thống theo dõi, Hệ thống "
    "điều phối (actors); :GiaoDienTaiXe (Boundary); :VanChuyenController "
    "(Controller); :ChuyenDi, :LoTrinh, :KienHang, :DonHang, :ViTriGPS, "
    ":BangChungGiaoHang, :SuCoGiaoThong (Entities); LogiFast DB / DAM "
    "(Database). Kịch bản thông điệp (100% Tiếng Việt, tự nhiên không tham số):"
)

make_table(
    headers=["#", "Người gửi", "Người nhận", "Thông điệp", "Bước UC", "Ghi chú"],
    rows=[
        ["1", "Tài xế", ":GiaoDienTaiXe", "Nhấn nút bắt đầu hành trình()", "B1", ""],
        ["2", ":GiaoDienTaiXe", ":VanChuyenController", "Yêu cầu bắt đầu ca()", "B1", ""],
        ["3", ":VanChuyenController", ":ChuyenDi", "Tạo chuyến đi()", "B1", ""],
        ["4", ":ChuyenDi", ":LoTrinh", "Khởi tạo lộ trình()", "B1", ""],
        ["5", "Tài xế", ":GiaoDienTaiXe", "Quét mã kiện hàng()", "B1", ""],
        ["6", ":VanChuyenController", ":KienHang", "Truy xuất thông tin kiện hàng()", "B1", ""],
        ["7", ":VanChuyenController", ":DonHang", "Báo cáo thông tin đơn khởi tạo()", "B1", ""],
        ["8", ":VanChuyenController", "Hệ thống điều phối", "Yêu cầu tính toán lộ trình tối ưu()", "B1", ""],
        ["9", ":VanChuyenController", ":LoTrinh", "Cập nhật dữ liệu lộ trình()", "B1", ""],
        ["10", "Hệ thống theo dõi", ":VanChuyenController", "Báo cáo Tọa độ()", "B2-B3 loop", ""],
        ["11", ":VanChuyenController", ":ViTriGPS", "Lưu vết định vị GPS()", "B2-B3", ""],
        ["12", ":VanChuyenController", ":LoTrinh", "Bổ sung tọa độ vào lộ trình()", "B2-B3", ""],
        ["13", ":VanChuyenController", "Khách hàng", "Gửi thông báo sắp đến nơi()", "B3", "[Khoảng cách < 200m]"],
        ["14", "Hệ thống theo dõi", ":VanChuyenController", "Báo động phát hiện tắc đường()", "B4 alt", ""],
        ["15", ":VanChuyenController", ":SuCoGiaoThong", "Ghi nhận sự cố()", "B4 alt", ""],
        ["16", ":VanChuyenController", ":LoTrinh", "Thay đổi trạng thái Đang tái tối ưu()", "B4 alt", ""],
        ["17", ":VanChuyenController", "Hệ thống điều phối", "Yêu cầu tính lại lộ trình mới()", "B4 alt", ""],
        ["18", ":VanChuyenController", ":GiaoDienTaiXe", "Cảnh báo đổi hướng di chuyển()", "B4 alt", ""],
        ["19", "Tài xế", ":GiaoDienTaiXe", "Nhấn nút đã đến nơi()", "B6", ""],
        ["20", "Tài xế", "Khách hàng", "Trao đổi vật lý()", "B6", "Hoạt động ngoại tuyến"],
        ["21", "Tài xế", ":GiaoDienTaiXe", "Xác nhận giao thành công()", "B7", ""],
        ["22", ":VanChuyenController", ":BangChungGiaoHang", "Lưu hình ảnh bằng chứng()", "B7", ""],
        ["23", ":VanChuyenController", ":DonHang", "Đánh dấu trạng thái Giao thành công()", "B7", ""],
        ["24", ":VanChuyenController", ":KienHang", "Đánh dấu Đã giao hiện tại()", "B7", ""],
        ["25", ":VanChuyenController", ":LoTrinh", "Khép lại tiến trình hiện tại()", "B7", ""],
        ["26", ":VanChuyenController", "LogiFast DB / DAM", "Lưu trữ biến động vào cơ sở dữ liệu()", "B7", ""],
    ],
    font_size=7,
)
spacer()
add_placeholder("Chèn sơ đồ tuần tự mức nghiệp vụ UC-03 — Project/Diagrams/UC03_SSD_business.drawio")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SƠ ĐỒ GIAO TIẾP MỨC THIẾT KẾ / NGHIỆP VỤ
# ══════════════════════════════════════════════════════════════════════════════

add_h1("5. Sơ đồ giao tiếp mức thiết kế / nghiệp vụ")

add_para(
    "Sơ đồ giao tiếp (Communication Diagram) ở mức nghiệp vụ tập trung vào khía "
    "cạnh kiến trúc: thay vì bố cục theo thời gian như sơ đồ tuần tự, nó sắp "
    "xếp các đối tượng tự do dạng mạng lưới và hiển thị rõ ràng các liên kết "
    "(Links) cấu trúc. Điều này tương ứng chặt chẽ với những gì đã phân tích "
    "trong Thẻ CRC: đối tượng nào \"biết\" đối tượng nào để gửi thông điệp."
)

add_h2("5.1. Kịch bản thông điệp — UC-03 (Business-Level CD)")

add_para(
    "Các móc nối liên kết (Links) dựa trên Thẻ CRC và Sơ đồ lớp: Tài xế liên "
    "kết với :GiaoDienTaiXe; :GiaoDienTaiXe gửi yêu cầu vào :VanChuyenController; "
    "Hệ thống theo dõi và Hệ thống điều phối kết nối vào :VanChuyenController; "
    ":VanChuyenController nắm giữ liên kết đến hầu hết các Entity; :ChuyenDi có "
    "trách nhiệm khởi tạo :LoTrinh theo thẻ CRC, nên có link điều khiển trực "
    "tiếp tới :LoTrinh mà không cần Controller can thiệp."
)

make_table(
    headers=["#", "Người gửi", "Người nhận", "Thông điệp", "Bước UC", "Ghi chú"],
    rows=[
        ["1", "Tài xế", ":GiaoDienTaiXe", "Nhấn nút bắt đầu hành trình()", "B1", ""],
        ["1.1", ":GiaoDienTaiXe", ":VanChuyenController", "Yêu cầu bắt đầu ca()", "B1", ""],
        ["1.1.1", ":VanChuyenController", ":ChuyenDi", "Tạo chuyến đi()", "B1", ""],
        ["1.1.2", ":ChuyenDi", ":LoTrinh", "Khởi tạo lộ trình()", "B1", ""],
        ["2", "Tài xế", ":GiaoDienTaiXe", "Quét mã kiện hàng()", "B1", ""],
        ["2.1", ":GiaoDienTaiXe", ":VanChuyenController", "Khởi tạo quá trình xử lý kiện hàng()", "B1", ""],
        ["2.1.1", ":VanChuyenController", ":KienHang", "Truy xuất thông tin kiện hàng()", "B1", ""],
        ["2.1.2", ":VanChuyenController", ":DonHang", "Báo cáo thông tin đơn khởi tạo()", "B1", ""],
        ["2.1.3", ":VanChuyenController", "Hệ thống điều phối", "Yêu cầu tính toán lộ trình tối ưu()", "B1", ""],
        ["2.1.4", ":VanChuyenController", ":LoTrinh", "Cập nhật dữ liệu lộ trình()", "B1", ""],
        ["3", "Hệ thống theo dõi", ":VanChuyenController", "Báo cáo Tọa độ()", "B2-B3", "loop"],
        ["3.1", ":VanChuyenController", ":ViTriGPS", "Lưu vết định vị GPS()", "B2-B3", ""],
        ["3.2", ":VanChuyenController", ":LoTrinh", "Bổ sung tọa độ vào lộ trình()", "B2-B3", ""],
        ["3.3", ":VanChuyenController", "Khách hàng", "Gửi thông báo sắp đến nơi()", "B3", "[Khoảng cách < 200m]"],
        ["4", "Hệ thống theo dõi", ":VanChuyenController", "Báo động phát hiện tắc đường()", "B4", "alt"],
        ["4.1", ":VanChuyenController", ":SuCoGiaoThong", "Ghi nhận sự cố()", "B4", ""],
        ["4.2", ":VanChuyenController", ":LoTrinh", "Thay đổi trạng thái Đang tái tối ưu()", "B4", ""],
        ["4.3", ":VanChuyenController", "Hệ thống điều phối", "Yêu cầu tính lại lộ trình mới()", "B4", ""],
        ["4.4", ":VanChuyenController", ":GiaoDienTaiXe", "Cảnh báo đổi hướng di chuyển()", "B4", ""],
        ["5", "Tài xế", ":GiaoDienTaiXe", "Nhấn nút đã đến nơi()", "B6", ""],
        ["6", "Tài xế", "Khách hàng", "Trao đổi vật lý()", "B6", "ngoại tuyến"],
        ["7", "Tài xế", ":GiaoDienTaiXe", "Xác nhận giao thành công()", "B7", ""],
        ["7.1", ":GiaoDienTaiXe", ":VanChuyenController", "Thông báo giao hàng thành công()", "B7", ""],
        ["7.1.1", ":VanChuyenController", ":BangChungGiaoHang", "Lưu hình ảnh bằng chứng()", "B7", ""],
        ["7.1.2", ":VanChuyenController", ":DonHang", "Đánh dấu trạng thái Giao thành công()", "B7", ""],
        ["7.1.3", ":VanChuyenController", ":KienHang", "Đánh dấu Đã giao hiện tại()", "B7", ""],
        ["7.1.4", ":VanChuyenController", ":LoTrinh", "Khép lại tiến trình hiện tại()", "B7", ""],
        ["7.1.5", ":VanChuyenController", "LogiFast DB / DAM", "Lưu trữ biến động vào CSDL()", "B7", ""],
    ],
    font_size=7,
)
spacer()
add_placeholder("Chèn sơ đồ giao tiếp mức nghiệp vụ UC-03 — Project/Diagrams/UC03_CommDiagram_business.drawio")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — MA TRẬN CRUD(E)
# ══════════════════════════════════════════════════════════════════════════════

add_h1("6. Ma trận CRUD(E)")

add_h2("6.1. Phương pháp xây dựng")

add_para(
    "Ma trận CRUD(E) hỗ trợ xác định các mối quan hệ giữa ca sử dụng và lớp "
    "lĩnh vực bằng cách gán nhãn các trường hợp tương tác: C (Create — tạo đối "
    "tượng mới), R (Read — tra cứu thông tin được lưu trong đối tượng), "
    "U (Update — cập nhật giá trị thuộc tính), D (Delete — xóa đối tượng), "
    "E (Execute — yêu cầu đối tượng thực hiện một hành động nghiệp vụ). "
    "Ô trống nghĩa là ca sử dụng đó không tương tác với lớp đó. "
    "Giá trị E được bổ sung trong bài giảng (slide 43): LoTrinh tại UC-03 "
    "là ô duy nhất có E — OR engine được yêu cầu thực thi (không chỉ đọc hay "
    "ghi) khi GPS tracking phát hiện lệch tuyến."
)

add_para(
    "Mỗi giá trị trong ô được chú thích bằng số bước UC: C(3) = tạo mới tại "
    "Bước 3 luồng chính; U(5-alt) = cập nhật tại Bước 5 luồng thay thế. "
    "Nguồn dữ liệu: đặc tả UC chi tiết (Chương III) và bảng thông điệp SSD "
    "(Mục 3 chương này). Quy tắc này đảm bảo tiêu chí I.19 và I.20: "
    "toàn bộ 5 UC × 22 lớp lĩnh vực có giá trị truy xuất đến từng bước UC."
)

add_h2("6.2. Ma trận CRUD(E) — LogiFast")

# CRUD(E) matrix: 5 UC rows × 22 class columns (including SuCoGiaoThong)
classes = [
    "KhachHang", "DonHang", "SanPham", "KienHang", "ChiTietKienHang",
    "PhuongThucThanhToan", "TaiXe", "PhieuPhanCong", "VungDiaChi", "DiemHieuSuat",
    "LoTrinh", "ViTriGPS", "ChuyenDi", "BangChungGiaoHang", "XacNhanGiaoHang",
    "ThongBao", "GiaoDich", "HoaDon", "ViDienTu", "SoCai", "PhiDichVu",
    "SuCoGiaoThong",
]

crud_data = {
    "UC-01": {
        "KhachHang": "R(1)", "DonHang": "C(3)", "SanPham": "R(2)",
        "KienHang": "C(3)", "ChiTietKienHang": "C(3)", "PhuongThucThanhToan": "R(2)",
    },
    "UC-02": {
        "DonHang": "R(1)\nU(5)", "TaiXe": "R(2)", "PhieuPhanCong": "C(3)",
        "VungDiaChi": "R(2)", "DiemHieuSuat": "R(2)\nU(5)",
    },
    "UC-03": {
        "DonHang": "U(3)\nU(6)", "KienHang": "U(2)\nU(6)", "TaiXe": "R(1)",
        "PhieuPhanCong": "R(1)", "LoTrinh": "C(3)\nE(4)", "ViTriGPS": "C(4)",
        "ChuyenDi": "C(1)", "BangChungGiaoHang": "C(6)", "SuCoGiaoThong": "C(4-alt)",
    },
    "UC-04": {
        "KhachHang": "R(3)", "DonHang": "U(4)", "TaiXe": "R(1)",
        "BangChungGiaoHang": "R(1)", "XacNhanGiaoHang": "C(2)", "ThongBao": "C(4)",
    },
    "UC-05": {
        "DonHang": "U(4)", "PhuongThucThanhToan": "R(1)", "TaiXe": "R(1)\nU(4)",
        "XacNhanGiaoHang": "R(1)", "GiaoDich": "C(3)", "HoaDon": "C(3)",
        "ViDienTu": "U(4)", "SoCai": "U(4)", "PhiDichVu": "R(1)",
    },
}

uc_list = ["UC-01", "UC-02", "UC-03", "UC-04", "UC-05"]

# Build header + rows
headers_crud = ["UC"] + classes
rows_crud = []
for uc in uc_list:
    row = [uc]
    for cls in classes:
        row.append(crud_data[uc].get(cls, ""))
    rows_crud.append(row)

tbl_crud = make_table(
    headers=headers_crud,
    rows=rows_crud,
    font_size=7,
)

# Set column widths: UC label = 1.1cm, class columns = 0.65cm each
for row in tbl_crud.rows:
    for ci, cell in enumerate(row.cells):
        if ci == 0:
            set_cell_width(cell, 620)   # UC label ~1.1cm
        else:
            set_cell_width(cell, 370)   # class col ~0.65cm

# Center-align content in class cells
for row in tbl_crud.rows:
    for ci, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            if ci > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

spacer()

add_h2("6.3. Giải thích các ô trọng tâm")

add_para(
    "Năm quan sát chính từ ma trận: "
    "(1) DonHang có mặt ở cả 5 UC — cột đầy nhất — khớp với vai trò trung tâm "
    "đã xác nhận qua sơ đồ máy trạng thái Mục 2.1; "
    "(2) ViTriGPS chỉ có C(4) ở UC-03, phản ánh đúng bản chất: chỉ quy trình "
    "vận chuyển mới ghi vết GPS thời gian thực; "
    "(3) GiaoDich, HoaDon, SoCai chỉ xuất hiện ở UC-05 — nhất quán với UC-05 "
    "là điểm kết toán cuối cùng của toàn bộ chu trình; "
    "(4) Không UC nào có D (Delete) — hệ thống giao hàng yêu cầu lưu trữ lịch "
    "sử toàn bộ để đối soát tài chính và xử lý khiếu nại; "
    "(5) LoTrinh có cả U và E ở UC-03: U khi cập nhật tuyến sau tái tối ưu, "
    "E khi OR engine được kích hoạt thực thi — đây là ô duy nhất trong ma trận "
    "có giá trị E, thể hiện tính chủ động của hệ thống LogiFast."
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(str(OUT))
print(f"✓ Done → {OUT}")
print(f"✓ Generated: Sections 1–6 (Giới thiệu, Máy trạng thái, SSD, SD/CD nghiệp vụ, CRUD(E))")
print(f"✓ Placeholders left for diagrams (insert manually)")
