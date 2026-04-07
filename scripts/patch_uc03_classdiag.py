"""
scripts/patch_uc03_classdiag.py
Performs three operations on Ch4_MoHinhHoaCauTruc.docx:

  1. Remove the duplicate noun-analysis block from 4.3.1 (elements [65]-[70])
  2. Replace the '[Chèn sơ đồ lớp lĩnh vực]' placeholder in 4.3.3 with:
       • Intro paragraph
       • Class + attribute scheme table (7 classes, analysis-phase: name-only attrs)
       • Association scheme table (6 associations with multiplicities)
       • Placeholder line for the .drawio file
  3. Apply a full black border to EVERY table in the document
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── low-level helpers ─────────────────────────────────────────────────────────

def gettxt(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{NS}}}t"))

def detach(obj):
    e = obj._element if hasattr(obj, "_element") else obj
    e.getparent().remove(e)
    return e

def make_para(doc, text="", style="Normal", bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
    return detach(p)

def make_mixed(doc, segments):
    """segments = [(text, bold), ...]"""
    p = doc.add_paragraph(style="Normal")
    for text, bold in segments:
        if text:
            r = p.add_run(text)
            r.bold = bold
    return detach(p)

def _border_el(val="single", sz="6", color="000000", space="0"):
    el = OxmlElement("w:tcBorder") if False else OxmlElement("w:top")  # placeholder
    return el

def _make_border(tag, val="single", sz="6", color="000000", space="0"):
    el = OxmlElement(tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:color"), color)
    el.set(qn("w:space"), space)
    return el

def apply_cell_borders(cell):
    """Apply black single border to all 4 sides of a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    else:
        # clear existing
        for child in list(tcBorders):
            tcBorders.remove(child)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tcBorders.append(_make_border(f"w:{side}"))

def apply_table_borders(tbl):
    """Apply borders to every cell in a table."""
    for row in tbl.rows:
        for cell in row.cells:
            apply_cell_borders(cell)

def make_table(doc, headers, rows):
    """Create a bordered table, return detached (tbl_elem, spacer_elem)."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = "TableNormal"
    except Exception:
        pass

    # header row
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)

    apply_table_borders(tbl)      # ← borders on new table

    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def insert_before(elems, ref_elem):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)


# ── load document ─────────────────────────────────────────────────────────────

doc = Document(DOC_PATH)
body = doc.element.body
children = list(body)


# ════════════════════════════════════════════════════════════════════════════════
# STEP 1 — Remove duplicate noun-analysis block from 4.3.1
# Elements: "Phân tích danh từ UC-03" paragraph, its intro, table, spacer,
#           summary paragraph, and trailing blank
# ════════════════════════════════════════════════════════════════════════════════

to_remove = []
for child in children:
    txt = gettxt(child)
    if "Phân tích danh từ UC-03" in txt:
        to_remove.append(child)
    elif to_remove and not any(
        m in txt for m in ["4.3.2", "Bảng đối tượng UC-03"]
    ):
        to_remove.append(child)
    elif "Bảng đối tượng UC-03" in txt:
        break   # stop here — keep this and everything after

for elem in to_remove:
    body.remove(elem)

print(f"Step 1: removed {len(to_remove)} duplicate noun elements")


# ════════════════════════════════════════════════════════════════════════════════
# STEP 2 — Replace placeholder in 4.3.3 with class diagram description + scheme
# ════════════════════════════════════════════════════════════════════════════════

# Reload children after removal
children = list(body)

# Find heading 4.3.3 and its placeholder paragraph
h433_elem = None
placeholder_elem = None
for i, child in enumerate(children):
    txt = gettxt(child)
    if "4.3.3" in txt:
        h433_elem = child
    elif h433_elem is not None and "Chèn sơ đồ lớp lĩnh vực" in txt:
        placeholder_elem = child
        break
    elif h433_elem is not None and "4.3.4" in txt:
        break   # safety stop

assert placeholder_elem is not None, "Could not find 4.3.3 placeholder"

# Build replacement elements
new_elems = []

# ── Intro ─────────────────────────────────────────────────────────────────────
new_elems.append(make_para(doc,
    "Sơ đồ lớp lĩnh vực UC-03 mô tả 7 lớp tham gia trực tiếp vào quy trình "
    "vận chuyển đơn hàng: TaiXe, DonHang, KienHang, LoTrinh, ViTriGPS, "
    "ChuyenDi và BangChungGiaoHang. Đây là ca sử dụng phức tạp nhất (13 bước) "
    "nên sơ đồ lớp có số liên kết nhiều nhất so với 4 UC còn lại. "
    "Mọi thuộc tính được ghi ở mức phân tích: chỉ tên, không có kiểu dữ liệu "
    "hay ký hiệu phạm vi truy cập."
))

# ── Class + attribute scheme table ───────────────────────────────────────────
new_elems.append(make_mixed(doc, [
    ("Sơ đồ 4.3.3 — Lược đồ lớp lĩnh vực UC-03 (phân tích):", True),
]))

cls_h = ["Lớp", "Thuộc tính (phân tích — chỉ tên)", "Ghi chú"]
cls_r = [
    ["TaiXe",
     "id, ten, soDienThoai, trangThai, viTriHienTai",
     "Kế thừa từ UC-02; xuất hiện như actor thực thi lộ trình"],
    ["DonHang",
     "id, trangThai, diaChiGiao, tongGiaTri, hinhThucThanhToan, thoiGianDat",
     "Kế thừa từ UC-01/02; trangThai thay đổi qua 3 giá trị trong UC-03"],
    ["KienHang",
     "id, maQR, khoiLuong, trangThai",
     "Đơn vị vật lý; maQR dùng để xác nhận lấy hàng tại kho"],
    ["LoTrinh",
     "id, diemXuatPhat, diemDen, khoangCach, /thoiGianDuKien, trangThai",
     "/thoiGianDuKien là thuộc tính suy diễn từ khoangCach"],
    ["ViTriGPS",
     "id, viDo, kinhDo, thoiGian",
     "Mỗi bản ghi là 1 snapshot GPS; 1 LoTrinh có nhiều ViTriGPS"],
    ["ChuyenDi",
     "id, thoiGianBatDau, thoiGianKetThuc",
     "Đại diện ca làm việc; bao gồm ≥1 LoTrinh"],
    ["BangChungGiaoHang",
     "id, tenNguoiNhan, hinhAnh, toaDo, thoiGianGiao",
     "Tạo sau khi giao thành công; dùng lại ở UC-04"],
]
t, s = make_table(doc, cls_h, cls_r)
new_elems.extend([t, s])

# ── Association scheme table ──────────────────────────────────────────────────
new_elems.append(make_mixed(doc, [
    ("Bảng quan hệ — UC-03:", True),
]))
new_elems.append(make_para(doc,
    "Mọi quan hệ đều có tên và cơ số tại cả hai đầu theo quy tắc phân tích I.7."
))

assoc_h = ["Từ lớp", "Tên quan hệ", "Đến lớp", "Cơ số (Từ → Đến)", "Loại"]
assoc_r = [
    ["ChuyenDi",          "bao gồm",       "LoTrinh",           "1 → 1..*",  "Association"],
    ["TaiXe",             "thực hiện",     "LoTrinh",           "1 → 0..*",  "Association"],
    ["LoTrinh",           "giao",          "DonHang",           "1..* → 1",  "Association"],
    ["LoTrinh",           "ghi nhận",      "ViTriGPS",          "1 → 1..*",  "Composition"],
    ["DonHang",           "chứa",          "KienHang",          "1 → 1..*",  "Composition"],
    ["DonHang",           "có bằng chứng", "BangChungGiaoHang", "1 → 0..1",  "Association"],
]
t, s = make_table(doc, assoc_h, assoc_r)
new_elems.extend([t, s])

# ── Drawio placeholder ────────────────────────────────────────────────────────
new_elems.append(make_mixed(doc, [
    ("[Chèn sơ đồ lớp lĩnh vực UC-03 — file: ", False),
    ("Project/Diagrams/UC03_ClassDiagram.drawio", True),
    ("]", False),
]))

# Remove old placeholder, insert new elements before it
parent = placeholder_elem.getparent()
insert_before(new_elems, placeholder_elem)
parent.remove(placeholder_elem)

print(f"Step 2: replaced 4.3.3 placeholder with {len(new_elems)} elements")


# ════════════════════════════════════════════════════════════════════════════════
# STEP 3 — Apply borders to ALL tables in the document
# ════════════════════════════════════════════════════════════════════════════════

count = 0
for tbl in doc.tables:
    apply_table_borders(tbl)
    count += 1

print(f"Step 3: applied borders to {count} tables")


# ── save ──────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Done. Saved to", DOC_PATH)
