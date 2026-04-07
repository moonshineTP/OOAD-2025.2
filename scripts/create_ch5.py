"""
Creates Project/Ch5_MoHinhHoaHanhVi.docx
Chapter V: Mô hình hóa hành vi — LogiFast delivery system
Owner scope: UC-03 (Phạm Gia Hưng); other UC SSDs are placeholders for teammates.

Run from repo root:
    python scripts/create_ch5.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("Project/Ch4_MoHinhHoaCauTruc.docx")   # style donor
OUT = Path("Project/Ch5_MoHinhHoaHanhVi.docx")

# ── Bootstrap: copy Ch4 for styles, then clear body ────────────────────────────
shutil.copy(SRC, OUT)
doc = Document(str(OUT))

body = doc.element.body
# Keep only the last child (sectPr — page setup), remove everything else
children = list(body)
sect_pr = body.find(qn("w:sectPr"))
for child in children:
    if child is not sect_pr:
        body.remove(child)

# ── Helper: borders ─────────────────────────────────────────────────────────────

def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def shade_cells(row, hex_color="D9D9D9"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for existing in tcPr.findall(qn("w:shd")):
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def set_cell_font(cell, size_pt, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size_pt)
            if bold:
                run.bold = True

def make_table(headers, rows, font_size=10, header_bold=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_borders(tbl)
    # Header
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        set_cell_font(tbl.rows[0].cells[i], font_size, bold=header_bold)
    shade_cells(tbl.rows[0])
    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            tbl.rows[ri + 1].cells[ci].text = str(val) if val else ""
            set_cell_font(tbl.rows[ri + 1].cells[ci], font_size)
    return tbl

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def add_placeholder(text):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    return p

def add_code(code_text):
    """Insert preformatted PlantUML block."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    doc.add_paragraph()

def spacer():
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════════

doc.add_paragraph("V. Mô hình hóa hành vi", style="Title")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 1 — GIỚI THIỆU
# ════════════════════════════════════════════════════════════════════════════════

add_h1("1. Giới thiệu")

add_para(
    "Sau khi mô hình cấu trúc (Chương IV) xác định những gì hệ thống ghi nhớ — "
    "các lớp lĩnh vực, thuộc tính và mối quan hệ tĩnh giữa chúng — mô hình hóa "
    "hành vi trả lời câu hỏi như thế nào các đối tượng đó biến đổi trạng thái và "
    "tương tác ra sao với nhau trong từng ca sử dụng. Theo bài giảng Ch04 "
    "(TS. Nguyễn Bá Ngọc), ba vấn đề cốt lõi cần mô hình hóa là: (1) sự biến đổi "
    "trạng thái của đối tượng theo tiến trình nghiệp vụ; (2) tương tác giữa tác "
    "nhân và hệ thống; (3) tương tác giữa các đối tượng nội bộ để đáp ứng hoạt "
    "động nghiệp vụ."
)

add_para(
    "Chương V trình bày bốn nội dung theo đúng cấu trúc bài giảng: sơ đồ máy "
    "trạng thái cho các đối tượng trọng tâm (Mục 2); sơ đồ tuần tự mức hệ thống "
    "— SSD — cho cả 5 ca sử dụng (Mục 3); sơ đồ giao tiếp mức hệ thống cho "
    "UC-03 (Mục 4); và ma trận CRUD(E) tổng hợp (Mục 5). Sơ đồ tuần tự mức "
    "nghiệp vụ và sơ đồ giao tiếp mức nghiệp vụ thuộc giai đoạn thiết kế và sẽ "
    "được trình bày ở Phần II."
)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SƠ ĐỒ MÁY TRẠNG THÁI
# ════════════════════════════════════════════════════════════════════════════════

add_h1("2. Sơ đồ máy trạng thái")

add_para(
    "Theo tiêu chí lựa chọn từ bài giảng (slide 5), đối tượng được đưa vào mô "
    "hình máy trạng thái phải thỏa đồng thời: (a) xuất hiện trong nhiều ca sử "
    "dụng, (b) có trạng thái liên quan đến công việc nghiệp vụ được biểu diễn "
    "qua thuộc tính rõ ràng, và (c) có sự thay đổi trạng thái được kích hoạt bởi "
    "các sự kiện nghiệp vụ xác định. Trong LogiFast, ba đối tượng thỏa điều kiện: "
    "DonHang (xuyên suốt 5 UC với thuộc tính trangThai), TaiXe (điều kiện bảo vệ "
    "phân công trong UC-02), và LoTrinh (tái tối ưu động bởi OR engine trong "
    "UC-03 — trạng thái duy nhất có do-activity liên tục trong toàn hệ thống)."
)

# Terminology table
make_table(
    headers=["Thuật ngữ", "Mô tả"],
    rows=[
        ["Trạng thái", "Một giai đoạn trong tiến trình; bộ giá trị thuộc tính thỏa điều kiện nhất định"],
        ["Bước chuyển (Transition)", "Từ trạng thái nguồn → trạng thái đích; nhãn: sự_kiện [điều_kiện] / hành_động"],
        ["Sự kiện kích hoạt", "Nguyên nhân dẫn đến chuyển trạng thái"],
        ["Điều kiện bảo vệ", "Chỉ chuyển nếu điều kiện [...] được đáp ứng"],
        ["Biểu thức hành vi", "Được thực hiện và hoàn thành trước khi chuyển sang trạng thái đích"],
        ["Hoạt động bên trong (do-activity)", "Chạy liên tục trong khi đối tượng ở trạng thái đó; hoàn thành khi kết thúc"],
        ["Trạng thái tổng hợp", "Chứa máy trạng thái con với trạng thái bắt đầu và kết thúc riêng"],
    ],
    font_size=10,
)
spacer()

# ── 2.1 DonHang ────────────────────────────────────────────────────────────────

add_h2("2.1. Sơ đồ máy trạng thái — DonHang")

add_para(
    "DonHang là đối tượng trọng tâm duy nhất tham gia vào cả 5 UC với thuộc tính "
    "trangThai biến đổi qua các giai đoạn nghiệp vụ rõ ràng. Trạng thái tại bất "
    "kỳ thời điểm nào xác định hành vi hệ thống: ví dụ, hệ thống từ chối phân "
    "công tài xế nếu DonHang không ở trạng thái \"Chờ phân công\", và từ chối "
    "kết toán COD nếu chưa đạt \"Giao thành công\". Đây là đối tượng ví dụ điển "
    "hình như Ví dụ 4.1b trong bài giảng (slide 8)."
)

add_para(
    "Sáu trạng thái chính theo luồng sự kiện happy path (kịch bản ORD-2026-001 "
    "ngày 26/03/2026): Sẵn sàng giao → Chờ phân công → Đã phân công → Đang giao "
    "→ Giao thành công → Hoàn tất thanh toán. Hai trạng thái ngoại lệ: "
    "Giao thất bại và Đã hủy. Mỗi bước chuyển trạng thái được liên kết tường "
    "minh với một thông điệp SSD (Mục 3) và một bước trong đặc tả UC tương ứng, "
    "đảm bảo tính truy xuất theo tiêu chí I.18."
)

make_table(
    headers=["Từ trạng thái", "Sự kiện", "Điều kiện", "Hành động", "Đến trạng thái", "UC / Bước", "Thông điệp SSD"],
    rows=[
        ["●", "—", "—", "tạo DonHang", "Sẵn sàng giao", "UC-01 B2", "datDonHang(...)"],
        ["Sẵn sàng giao", "OrderCreated", "[đơn hợp lệ]", "gán maDonHang", "Chờ phân công", "UC-01 B3", "xacNhanDonHang(maDonHang)"],
        ["Chờ phân công", "AssignmentStarted", "[có tài xế phù hợp]", "tạo PhieuPhanCong", "Đã phân công", "UC-02 B3", "phanCongTaiXe(maDonHang, maTaiXe)"],
        ["Chờ phân công", "NoDriverFound", "[hết thời gian chờ]", "ghi log", "Giao thất bại", "UC-02 alt", "—"],
        ["Đã phân công", "DriverDeparted", "—", "cập nhật trangThai", "Đang giao", "UC-03 B3", "batDauGiaoHang(maDonHang)"],
        ["Đang giao", "DeliveryAttemptFailed", "[không liên lạc được]", "—", "Giao thất bại", "UC-03 alt", "—"],
        ["Đang giao", "PackageHandedOver", "—", "tạo BangChungGiaoHang", "Giao thành công", "UC-03 B5", "xacNhanDaDenDiemGiao(maDonHang)"],
        ["Giao thành công", "RecipientConfirmed", "—", "cập nhật trangThai", "Đã xác nhận", "UC-04 B4", "xacNhanGiaoHang(maXacNhan)"],
        ["Đã xác nhận", "PaymentSettled", "—", "cập nhật trangThai", "Hoàn tất thanh toán", "UC-05 B4", "ketToanCOD(maDonHang, soTien)"],
        ["Hoàn tất thanh toán", "—", "—", "—", "⊙", "—", "—"],
        ["Giao thất bại", "—", "—", "—", "⊙", "—", "—"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ máy trạng thái DonHang — Project/Diagrams/DonHang_StateMachine.puml")

# ── 2.2 TaiXe ──────────────────────────────────────────────────────────────────

add_h2("2.2. Sơ đồ máy trạng thái — TaiXe")

add_para(
    "TaiXe có vòng đời theo ca làm việc. Trạng thái của tài xế là điều kiện bảo "
    "vệ quan trọng trong UC-02: hệ thống phân công chỉ xét các tài xế đang ở "
    "trạng thái \"Sẵn sàng nhận đơn\". Ngoài ra, trạng thái \"Đang giao hàng\" "
    "giải thích tại sao cùng một tài xế không thể nhận thêm đơn mới trong UC-02 "
    "khi đang thực hiện UC-03. Vòng lặp Sẵn sàng nhận đơn ↔ Đang giao hàng có "
    "thể lặp lại nhiều lần trong một ca, tương ứng với số đơn hàng được giao."
)

make_table(
    headers=["Từ trạng thái", "Sự kiện", "Điều kiện", "Hành động", "Đến trạng thái", "UC / Bước", "Thông điệp SSD"],
    rows=[
        ["●", "—", "—", "—", "Ngoài ca", "—", "—"],
        ["Ngoài ca", "ShiftStarted", "—", "tạo ChuyenDi", "Sẵn sàng nhận đơn", "UC-03 B1", "xacNhanBatDauCa(maTaiXe)"],
        ["Sẵn sàng nhận đơn", "AssignmentReceived", "—", "cập nhật trangThai", "Đang giao hàng", "UC-02 B3", "phanCongTaiXe(...)"],
        ["Đang giao hàng", "DeliveryCompleted", "—", "cập nhật trangThai", "Sẵn sàng nhận đơn", "UC-03 B5", "xacNhanDaDenDiemGiao(...)"],
        ["Sẵn sàng nhận đơn", "ShiftEnded", "—", "ghi thoiGianKetThuc", "Kết thúc ca", "UC-03 (end)", "—"],
        ["Kết thúc ca", "—", "—", "—", "⊙", "—", "—"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ máy trạng thái TaiXe — Project/Diagrams/TaiXe_StateMachine.puml")

# ── 2.3 LoTrinh ────────────────────────────────────────────────────────────────

add_h2("2.3. Sơ đồ máy trạng thái — LoTrinh")

add_para(
    "LoTrinh là đối tượng kỹ thuật phức tạp nhất trong UC-03. Không giống các "
    "lớp khác, LoTrinh không có vòng đời tuyến tính: nó có thể bị tái tính toán "
    "(recalculated) bởi OR engine khi phát sinh sự kiện bất thường — tắc đường "
    "được phát hiện bởi Hệ thống theo dõi, hoặc GPS tài xế lệch khỏi tuyến đã "
    "tính hơn 500m, hoặc SuCoGiaoThong được ghi nhận. Trạng thái "
    "\"Đang tái tối ưu\" là trạng thái duy nhất trong toàn bộ mô hình LogiFast "
    "có hoạt động nội bộ liên tục (do / OR engine recalculates route), thể hiện "
    "tính chủ động của hệ thống."
)

add_para(
    "Điều kiện bảo vệ [GPS lệch >500m] được cấp trực tiếp từ thông điệp "
    "capNhatViTriGPS(viDo, kinhDo) trong vòng lặp \"loop [mỗi 30 giây]\" của SSD "
    "UC-03 (Mục 3.3). Đây là liên kết I.18: chuyển trạng thái Đang thực hiện → "
    "Đang tái tối ưu có thể truy xuất về cả bước UC-03 Bước 4 lẫn thông điệp SSD "
    "capNhatViTriGPS. Thuộc tính /thoiGianDuKien được tái tính mỗi lần OR hoàn "
    "thành, chứng minh thoiGianDuKien là thuộc tính suy diễn động (derived), "
    "không tĩnh — ký hiệu / theo quy ước UML."
)

make_table(
    headers=["Từ trạng thái", "Sự kiện", "Điều kiện", "Hành động", "Đến trạng thái", "UC / Bước", "Thông điệp SSD"],
    rows=[
        ["●", "—", "—", "OR engine tính tuyến", "Đã tính toán", "UC-03 B3", "batDauGiaoHang(maDonHang)"],
        ["Đã tính toán", "DriverDeparted", "[tài xế xác nhận]", "ghi thoiGianBatDau", "Đang thực hiện", "UC-03 B3", "hienThiLoTrinh(maLoTrinh, khoangCach, ETA)"],
        ["Đang thực hiện", "RouteDeviated", "[GPS lệch >500m OR SuCoGiaoThong]", "kích hoạt OR engine", "Đang tái tối ưu", "UC-03 B4", "capNhatViTriGPS(viDo, kinhDo)"],
        ["Đang tái tối ưu", "OptimizationDone", "[tuyến mới tốt hơn ≥10%]", "cập nhật khoangCach, /ETA", "Đang thực hiện", "UC-03 B4 loop", "—"],
        ["Đang tái tối ưu", "OptimizationDone", "[tuyến cũ vẫn tối ưu]", "giữ nguyên", "Đang thực hiện", "UC-03 B4 loop", "—"],
        ["Đang thực hiện", "PackageHandedOver", "[GPS ±200m điểm đến]", "ghi thoiGianKetThuc", "Hoàn tất", "UC-03 B5", "xacNhanDaDenDiemGiao(...)"],
        ["Đang thực hiện", "OrderCancelled", "—", "ghi lý do", "Bị hủy", "UC-03 alt", "—"],
        ["Hoàn tất", "—", "—", "—", "⊙", "—", "—"],
        ["Bị hủy", "—", "—", "—", "⊙", "—", "—"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ máy trạng thái LoTrinh — Project/Diagrams/LoTrinh_StateMachine.puml")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SƠ ĐỒ TUẦN TỰ MỨC HỆ THỐNG (SSD)
# ════════════════════════════════════════════════════════════════════════════════

add_h1("3. Sơ đồ tuần tự mức hệ thống (SSD)")

add_para(
    "Sơ đồ tuần tự mức hệ thống (System Sequence Diagram — SSD) có cấu trúc cố "
    "định: chỉ gồm tác nhân và hệ thống (\":Hệ thống LogiFast\"). Các thông điệp "
    "tương ứng với các hoạt động nghiệp vụ ở mức khái quát cao — mỗi thông điệp "
    "đại diện cho một bước trong đặc tả UC. SSD là đầu vào trực tiếp cho bước "
    "thiết kế: mỗi thông điệp gửi đến hệ thống sẽ trở thành một hợp đồng thông "
    "điệp (system operation contract) ở Phần II. Điều này đảm bảo tính nhất quán "
    "giữa mô hình phân tích và mô hình thiết kế."
)

add_para(
    "Quy tắc ký hiệu (slides 17–18, 23): tên đường sống viết dạng "
    "tênĐốiTượng:TênLớp, gạch chân, căn giữa. SSD chỉ có đúng hai đường sống: "
    "tác nhân (actor) và \":Hệ thống LogiFast\". Mô tả thông điệp theo cú pháp: "
    "[biểuThứcLôGic] tênThôngĐiệp(danhSáchThamSố). Điều kiện bảo vệ trong [...] "
    "xác định khi nào thông điệp được gửi. Khung kết hợp (combined fragments): "
    "loop (lặp), opt (tùy chọn), alt (rẽ nhánh), break (dừng vòng lặp)."
)

# ── 3.1 UC-01 placeholder ──────────────────────────────────────────────────────

add_h2("3.1. SSD — UC-01: Đặt đơn hàng")
add_para("Tác nhân: Khách hàng. UC-01 mô tả luồng khách hàng xác nhận giỏ hàng, "
         "nhập địa chỉ giao, chọn hình thức thanh toán và đặt đơn.")
add_placeholder("Chèn SSD UC-01 — Project/Diagrams/UC01_SSD.puml  |  Owner: Trương Văn Hồng")

# ── 3.2 UC-02 placeholder ──────────────────────────────────────────────────────

add_h2("3.2. SSD — UC-02: Phân công giao hàng")
add_para("Tác nhân: Hệ thống phân công (tự động). UC-02 mô tả quy trình hệ thống "
         "tìm tài xế phù hợp, tạo phiếu phân công và thông báo tài xế chấp nhận.")
add_placeholder("Chèn SSD UC-02 — Project/Diagrams/UC02_SSD.puml  |  Owner: Nguyễn Quý Duy")

# ── 3.3 UC-03 — FULL ───────────────────────────────────────────────────────────

add_h2("3.3. SSD — UC-03: Vận chuyển đơn hàng")

add_para(
    "Tác nhân: Tài xế giao hàng. UC-03 là ca sử dụng phức tạp nhất với 13 bước "
    "nghiệp vụ trải qua ba giai đoạn: lấy hàng tại kho (B1–B2), theo dõi hành "
    "trình (B3–B4), và bàn giao tại điểm đến (B5). SSD gồm 5 thông điệp gửi và "
    "5 thông điệp trả về, với một khung lặp loop bao phủ bước B4 (GPS tracking "
    "mỗi 30 giây). Tái định tuyến do SuCoGiaoThong xảy ra nội bộ trong hệ thống "
    "— không tạo thêm thông điệp mới ra tác nhân."
)

make_table(
    headers=["#", "Hướng", "Thông điệp (tham số)", "Bước UC", "Ghi chú"],
    rows=[
        ["1", "TX → :Hệ thống", "xacNhanBatDauCa(maTaiXe)", "B1 — bắt đầu ca", "Tạo ChuyenDi mới"],
        ["↩", ":Hệ thống → TX", "xacNhanChuyenDi(maChuyenDi, danhSachDon)", "—", "Trả danh sách đơn cần giao trong ca"],
        ["2", "TX → :Hệ thống", "quetMaQR(maQR)", "B2 — xác nhận kiện hàng", "Quét mã QR trên kiện hàng tại kho"],
        ["↩", ":Hệ thống → TX", "xacNhanKienHang(tenSanPham, diaChiGiao)", "—", "Xác định KienHang hợp lệ"],
        ["3", "TX → :Hệ thống", "batDauGiaoHang(maDonHang)", "B3 — xuất phát", "DonHang.trangThai → Đang giao"],
        ["↩", ":Hệ thống → TX", "hienThiLoTrinh(maLoTrinh, khoangCach, thoiGianDuKien)", "—", "Tính LoTrinh tối ưu, trả ETA"],
        ["4 (loop)", "TX → :Hệ thống", "capNhatViTriGPS(viDo, kinhDo)", "B4 — tracking [mỗi 30 giây]", "Ghi ViTriGPS; kích hoạt OR nếu lệch >500m"],
        ["↩", ":Hệ thống → TX", "ghiNhanThanhCong()", "—", "ETA hiệu chỉnh liên tục"],
        ["5", "TX → :Hệ thống", "xacNhanDaDenDiemGiao(maDonHang)", "B5 — đến điểm giao", "DonHang.trangThai → Chờ xác nhận; kích hoạt UC-04"],
        ["↩", ":Hệ thống → TX", "xacNhanCoMat(maThongBao)", "—", "Gửi ThongBao cho KhachHang — HandoverStarted"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn SSD UC-03 — Project/Diagrams/UC03_SSD.puml")

# ── 3.4 UC-04 placeholder ──────────────────────────────────────────────────────

add_h2("3.4. SSD — UC-04: Xác nhận giao hàng")
add_para("Tác nhân: Shipper (TaiXe). UC-04 mô tả luồng Shipper mở phiếu, yêu cầu "
         "OTP, nhận OTP từ khách hàng và xác nhận giao hàng thành công.")
add_placeholder("Chèn SSD UC-04 — Project/Diagrams/UC04_SSD.puml  |  Owner: Đinh Việt Hùng")

# ── 3.5 UC-05 placeholder ──────────────────────────────────────────────────────

add_h2("3.5. SSD — UC-05: Giao hàng hoàn tất / Thanh toán")
add_para("Tác nhân: Shipper (TaiXe). UC-05 mô tả luồng Shipper xem thông tin kết "
         "toán, xác nhận nhận tiền COD và thực hiện kết toán với hệ thống.")
add_placeholder("Chèn SSD UC-05 — Project/Diagrams/UC05_SSD.puml  |  Owner: Nguyễn Ngọc Toàn")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 4 — SƠ ĐỒ GIAO TIẾP
# ════════════════════════════════════════════════════════════════════════════════

add_h1("4. Sơ đồ tương tác")

add_para(
    "Sau khi SSD (Mục 3) xác định ranh giới hệ thống — những gì tác nhân gửi và "
    "nhận — mục này đi sâu vào hai loại sơ đồ tương tác bổ sung cho nhau: "
    "(a) sơ đồ giao tiếp mức hệ thống (§4.1) giúp đối chiếu với SSD dưới ký hiệu "
    "tự do (không bị gò theo trục thời gian), và (b) sơ đồ tuần tự mức nghiệp vụ "
    "(§4.2) chuyển sang bên trong hệ thống — chỉ rõ các đối tượng lĩnh vực nào "
    "thực sự cộng tác để hoàn thành từng bước UC-03. Cả hai cùng truy xuất về "
    "thẻ CRC (Ch IV §4.3.4): một lớp chỉ gửi thông điệp đến lớp nằm trong danh "
    "sách cộng tác viên của nó."
)

# ── 4.1 System-level communication diagram ────────────────────────────────────

add_h2("4.1. Sơ đồ giao tiếp mức hệ thống — UC-03")

add_para(
    "Sơ đồ giao tiếp (Communication Diagram) biểu diễn cùng nội dung với SSD "
    "nhưng với bố cục tự do: thứ tự thông điệp được đánh số theo biểu thức thứ "
    "tự (sequence expression) — \"1: m1()\", \"1.1: m2()\" (thông điệp con kích "
    "hoạt bởi 1), \"1a / 1b\" (song song). Tại mức hệ thống, sơ đồ vẫn chỉ gồm "
    "tác nhân và \":Hệ thống LogiFast\". Mỗi thông điệp tương ứng hoàn toàn với "
    "bảng thông điệp SSD Mục 3.3."
)

make_table(
    headers=["Số thứ tự", "Chiều", "Thông điệp (tham số)", "Bước UC", "Ghi chú"],
    rows=[
        ["1", "TX → :Hệ thống", "xacNhanBatDauCa(maTaiXe)", "B1", "Khởi tạo ca giao hàng"],
        ["1.1", ":Hệ thống → TX", "xacNhanChuyenDi(maChuyenDi, danhSachDon)", "—", "Phản hồi của 1"],
        ["2", "TX → :Hệ thống", "quetMaQR(maQR)", "B2", "Xác nhận kiện hàng tại kho"],
        ["2.1", ":Hệ thống → TX", "xacNhanKienHang(tenSanPham, diaChiGiao)", "—", "Phản hồi của 2"],
        ["3", "TX → :Hệ thống", "batDauGiaoHang(maDonHang)", "B3", "Xuất phát, khởi tạo LoTrinh"],
        ["3.1", ":Hệ thống → TX", "hienThiLoTrinh(maLoTrinh, khoangCach, ETA)", "—", "Phản hồi của 3"],
        ["4 [loop]", "TX → :Hệ thống", "capNhatViTriGPS(viDo, kinhDo)", "B4", "Lặp mỗi 30 giây"],
        ["4.1", ":Hệ thống → TX", "ghiNhanThanhCong()", "—", "Phản hồi của 4"],
        ["5", "TX → :Hệ thống", "xacNhanDaDenDiemGiao(maDonHang)", "B5", "Kết thúc UC-03"],
        ["5.1", ":Hệ thống → TX", "xacNhanCoMat(maThongBao)", "—", "Phản hồi của 5; kích hoạt UC-04"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ giao tiếp mức hệ thống UC-03 — Project/Diagrams/UC03_CommDiagram.puml")

# ── 4.2 Domain-level sequence diagram ─────────────────────────────────────────

add_h2("4.2. Sơ đồ tuần tự mức nghiệp vụ — UC-03")

add_para(
    "Sơ đồ tuần tự mức nghiệp vụ (business-level sequence diagram) chuyển góc "
    "nhìn từ ranh giới hệ thống vào bên trong: thay vì chỉ hai đường sống "
    "(actor + :Hệ thống), sơ đồ này triển khai các đối tượng lĩnh vực cụ thể "
    "tham gia vào UC-03 — cd:ChuyenDi, lt:LoTrinh, dh:DonHang, gps:ViTriGPS, "
    "sg:SuCoGiaoThong — và chỉ rõ thông điệp nào được truyền giữa chúng. "
    "Mỗi mũi tên có thể truy xuất về một quan hệ trong sơ đồ lớp UC-03 (Ch IV "
    "§4.3.2) và một cộng tác viên trong thẻ CRC tương ứng (Ch IV §4.3.4). "
    "Đây là cơ sở trực tiếp để thiết kế các hợp đồng thông điệp chi tiết "
    "ở Phần II."
)

make_table(
    headers=["#", "Người gửi", "Người nhận", "Thông điệp (tham số)", "Bước UC", "CRC grounding"],
    rows=[
        ["1", "tx:TaiXe", "cd:ChuyenDi", "tao(maTaiXe, ngayLamViec)", "B1", "TaiXe → ChuyenDi: thực hiện"],
        ["1.1", "cd:ChuyenDi", "lt:LoTrinh", "taoLoTrinh(danhSachDon)", "B1", "ChuyenDi → LoTrinh: bao gồm"],
        ["1.2", "lt:LoTrinh", "dh:DonHang", "layThongTin(maDon)", "B1", "LoTrinh → DonHang: giao"],
        ["2", "tx:TaiXe", "lt:LoTrinh", "xacNhanKienHang(maQR)", "B2", "TaiXe → LoTrinh: thực hiện"],
        ["2.1", "lt:LoTrinh", "dh:DonHang", "kiemTraMaQR(maQR)", "B2", "LoTrinh → DonHang: giao"],
        ["3", "tx:TaiXe", "lt:LoTrinh", "batDau()", "B3", "TaiXe → LoTrinh: thực hiện"],
        ["3.1", "lt:LoTrinh", "dh:DonHang", "setTrangThai(\"Đang giao\")", "B3", "LoTrinh → DonHang: giao"],
        ["4 loop", "tx:TaiXe", "gps:ViTriGPS", "ghi(viDo, kinhDo, ts)", "B4", "LoTrinh → ViTriGPS: ghi nhận"],
        ["4.1", "gps:ViTriGPS", "lt:LoTrinh", "capNhat(viTriMoi)", "B4", "ViTriGPS → LoTrinh"],
        ["4.2 alt", "lt:LoTrinh", "sg:SuCoGiaoThong", "tao(loaiSuCo, toaDo)", "B4 alt", "LoTrinh → SuCoGiaoThong"],
        ["4.3 alt", "lt:LoTrinh", "lt:LoTrinh", "taiToiUuTuyen()", "B4 alt", "do-activity: OR engine"],
        ["5", "tx:TaiXe", "lt:LoTrinh", "xacNhanDenDiemGiao(maDon)", "B5", "TaiXe → LoTrinh: thực hiện"],
        ["5.1", "lt:LoTrinh", "dh:DonHang", "setTrangThai(\"Giao thành công\")", "B5", "LoTrinh → DonHang: giao"],
    ],
    font_size=9,
)
spacer()
add_placeholder("Chèn sơ đồ tuần tự mức nghiệp vụ UC-03 — Project/Diagrams/UC03_DomainSequence.puml")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 5 — MA TRẬN CRUD(E)
# ════════════════════════════════════════════════════════════════════════════════

add_h1("5. Ma trận CRUD(E)")

add_h2("5.1. Phương pháp xây dựng")

add_para(
    "Ma trận CRUD(E) hỗ trợ xác định các mối quan hệ giữa ca sử dụng và lớp "
    "lĩnh vực bằng cách gán nhãn các trường hợp tương tác: C (Create — tạo đối "
    "tượng mới), R (Read — tra cứu thông tin được lưu trong đối tượng), "
    "U (Update — cập nhật giá trị thuộc tính), D (Delete — xóa đối tượng), "
    "E (Execute — yêu cầu đối tượng thực hiện một hành động nghiệp vụ). "
    "Ô trống nghĩa là ca sử dụng đó không tương tác với lớp đó. "
    "Giá trị E được bổ sung trong bài giảng (slide 43): LoTrinh tại UC-03 "
    "là ô duy nhất có E — OR engine được yêu cầu thực thi (không chỉ đọc hay "
    "ghi) khi GPS tracking phát hiện lệch tuyến."
)

add_para(
    "Mỗi giá trị trong ô được chú thích bằng số bước UC: C(3) = tạo mới tại "
    "Bước 3 luồng chính; U(5-alt) = cập nhật tại Bước 5 luồng thay thế. "
    "Nguồn dữ liệu: đặc tả UC chi tiết (Chương III) và bảng thông điệp SSD "
    "(Mục 3 chương này). Quy tắc này đảm bảo tiêu chí I.19 và I.20: "
    "toàn bộ 5 UC × 22 lớp lĩnh vực có giá trị truy xuất đến từng bước UC."
)

add_h2("5.2. Ma trận CRUD(E) — LogiFast")

# CRUD(E) matrix: 5 UC rows × 22 class columns
classes = [
    "KhachHang", "DonHang", "SanPham", "KienHang", "ChiTietKienHang",
    "PhuongThucThanhToan", "TaiXe", "PhieuPhanCong", "VungDiaChi", "DiemHieuSuat",
    "LoTrinh", "ViTriGPS", "ChuyenDi", "BangChungGiaoHang", "XacNhanGiaoHang",
    "ThongBao", "GiaoDich", "HoaDon", "ViDienTu", "SoCai", "PhiDichVu",
    "SuCoGiaoThong",
]

crud_data = {
    "UC-01": {
        "KhachHang": "R(1)", "DonHang": "C(3)", "SanPham": "R(2)",
        "KienHang": "C(3)", "ChiTietKienHang": "C(3)", "PhuongThucThanhToan": "R(2)",
    },
    "UC-02": {
        "DonHang": "R(1)\nU(5)", "TaiXe": "R(2)", "PhieuPhanCong": "C(3)",
        "VungDiaChi": "R(2)", "DiemHieuSuat": "R(2)\nU(5)",
    },
    "UC-03": {
        "DonHang": "U(3)\nU(5)", "KienHang": "U(2)\nU(5)", "TaiXe": "R(1)",
        "PhieuPhanCong": "R(1)", "LoTrinh": "C(3)\nE(4)", "ViTriGPS": "C(4)",
        "ChuyenDi": "C(1)", "BangChungGiaoHang": "—", "SuCoGiaoThong": "R(4)",
    },
    "UC-04": {
        "KhachHang": "R(3)", "DonHang": "U(4)", "TaiXe": "R(1)",
        "BangChungGiaoHang": "R(1)", "XacNhanGiaoHang": "C(2)", "ThongBao": "C(4)",
    },
    "UC-05": {
        "DonHang": "U(4)", "PhuongThucThanhToan": "R(1)", "TaiXe": "R(1)\nU(4)",
        "XacNhanGiaoHang": "R(1)", "GiaoDich": "C(3)", "HoaDon": "C(3)",
        "ViDienTu": "U(4)", "SoCai": "U(4)", "PhiDichVu": "R(1)",
    },
}

uc_list = ["UC-01", "UC-02", "UC-03", "UC-04", "UC-05"]

# Build header + rows
headers_crud = ["UC"] + classes
rows_crud = []
for uc in uc_list:
    row = [uc]
    for cls in classes:
        row.append(crud_data[uc].get(cls, ""))
    rows_crud.append(row)

tbl_crud = make_table(
    headers=headers_crud,
    rows=rows_crud,
    font_size=7,
)

# Narrow column widths: UC label = 1.1cm, class columns = 0.65cm each
def set_tw(cell, twips):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

for row in tbl_crud.rows:
    for ci, cell in enumerate(row.cells):
        if ci == 0:
            set_tw(cell, 620)   # UC label ~1.1cm
        else:
            set_tw(cell, 370)   # class col ~0.65cm

# Center-align content in class cells
for row in tbl_crud.rows:
    for ci, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            if ci > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

spacer()

add_h2("5.3. Giải thích các ô trọng tâm")

add_para(
    "Năm quan sát chính từ ma trận: "
    "(1) DonHang có mặt ở cả 5 UC — cột đầy nhất — khớp với vai trò trung tâm "
    "đã xác nhận qua sơ đồ máy trạng thái Mục 2.1; "
    "(2) ViTriGPS chỉ có C(4) ở UC-03, phản ánh đúng bản chất: chỉ quy trình "
    "vận chuyển mới ghi vết GPS thời gian thực; "
    "(3) GiaoDich, HoaDon, SoCai chỉ xuất hiện ở UC-05 — nhất quán với UC-05 "
    "là điểm kết toán cuối cùng của toàn bộ chu trình; "
    "(4) Không UC nào có D (Delete) — hệ thống giao hàng yêu cầu lưu trữ lịch "
    "sử toàn bộ để đối soát tài chính và xử lý khiếu nại; "
    "(5) LoTrinh có cả U và E ở UC-03: U khi cập nhật tuyến sau tái tối ưu, "
    "E khi OR engine được kích hoạt thực thi — đây là ô duy nhất trong ma trận "
    "có giá trị E, thể hiện tính chủ động của hệ thống LogiFast."
)

# ════════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════════

doc.save(str(OUT))
print(f"Done → {OUT}")
