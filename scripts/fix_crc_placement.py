"""
scripts/fix_crc_placement.py
Moves the 5 CRC tables that were wrongly inserted into UC-01 (4.1.5)
to the correct location: UC-03 (4.3.5), before the still-existing placeholder.
Also fixes the UC-01 4.1.5 section by putting back the original placeholder.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document(DOC_PATH)
body = doc.element.body

def gettxt(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{NS}}}t"))

def detach(elem):
    elem.getparent().remove(elem)
    return elem

def insert_before(elems, ref_elem):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)

children = list(body)

# ── 1. Find boundaries of UC-01 section (4.1.5. Thẻ CRC) ─────────────────────
# Identify: the heading, the intro+cards that don't belong, and the next heading
uc01_crc_heading_idx = None
uc02_heading_idx = None

for i, child in enumerate(children):
    txt = gettxt(child)
    tag = child.tag.split("}")[1]
    style_el = child.find(f".//{{{NS}}}pStyle")
    sval = style_el.get(f"{{{NS}}}val") if style_el is not None else ""

    if "4.1.5" in txt and "CRC" in txt:
        uc01_crc_heading_idx = i
    if "4.2." in txt and uc01_crc_heading_idx is not None and i > uc01_crc_heading_idx:
        uc02_heading_idx = i
        break

print(f"UC-01 CRC heading at [{uc01_crc_heading_idx}], UC-02 at [{uc02_heading_idx}]")

# Elements in UC-01 CRC section (between heading and 4.2 heading, exclusive)
uc01_crc_section = children[uc01_crc_heading_idx + 1 : uc02_heading_idx]

# The wrongly-placed elements are: intro para, blank, 5×(table + spacer) = 12
# Identify: intro starts with "Thẻ CRC (Class", blank, then tables of CRC cards
# Keep: nothing — remove ALL wrong elements, restore just a placeholder
to_remove = []
for elem in uc01_crc_section:
    txt = gettxt(elem)
    tag = elem.tag.split("}")[1]
    # Remove intro, blank, tables beginning with "Tên lớp" (CRC templates)
    if txt.strip() == "" or "Thẻ CRC (Class" in txt or "Tên lớp" in txt:
        to_remove.append(elem)

print(f"Removing {len(to_remove)} elements from UC-01 CRC section")
for elem in to_remove:
    body.remove(elem)

# Restore UC-01 placeholder
children = list(body)
uc01_crc_heading_idx = None
for i, child in enumerate(children):
    if "4.1.5" in gettxt(child) and "CRC" in gettxt(child):
        uc01_crc_heading_idx = i; break

restore_p = OxmlElement("w:p")
restore_r = OxmlElement("w:r")
restore_t = OxmlElement("w:t")
restore_t.text = "[Chèn thẻ CRC các lớp chính]"
restore_r.append(restore_t)
restore_p.append(restore_r)

next_after_heading = children[uc01_crc_heading_idx + 1]
body.insert(list(body).index(next_after_heading), restore_p)
print(f"Restored placeholder in UC-01 4.1.5")

# ── 2. Now build and insert CRC cards at UC-03 4.3.5 ─────────────────────────

def apply_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
    else:
        for child in list(tcBorders): tcBorders.remove(child)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000"); el.set(qn("w:space"), "0")
        tcBorders.append(el)

def apply_table_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            apply_cell_borders(cell)

def detach2(obj):
    e = obj._element if hasattr(obj, "_element") else obj
    e.getparent().remove(e)
    return e

def make_para2(text="", bold=False):
    p = doc.add_paragraph(style="Normal")
    if text:
        r = p.add_run(text); r.bold = bold
    return detach2(p)

def make_crc_table(tbl_name, tbl_id, layer, desc, attrs, responsibilities, partners):
    tbl = doc.add_table(rows=7, cols=2)
    try: tbl.style = "TableNormal"
    except Exception: pass
    data = [
        ("Tên lớp", tbl_name), ("ID", tbl_id), ("Tầng kiến trúc", layer),
        ("Mô tả", desc), ("Các thuộc tính", attrs),
        ("Các trách nhiệm", responsibilities), ("Các đối tác", partners),
    ]
    for ri, (label, value) in enumerate(data):
        cells = tbl.rows[ri].cells
        cells[0].text = label
        for para in cells[0].paragraphs:
            for run in para.runs: run.bold = True
        cells[1].text = value
    apply_table_borders(tbl)
    tbl_elem = detach2(tbl)
    spacer = detach2(doc.add_paragraph())
    return tbl_elem, spacer

# Find the UC-03 CRC placeholder specifically (after 4.3.5 heading)
children = list(body)
uc03_crc_placeholder = None
in_uc03_crc = False
for child in children:
    txt = gettxt(child)
    if "4.3.5" in txt and "CRC" in txt:
        in_uc03_crc = True
    elif in_uc03_crc and "Chèn thẻ CRC" in txt:
        uc03_crc_placeholder = child
        break
    elif in_uc03_crc and "4.4." in txt:
        break  # safety

assert uc03_crc_placeholder is not None, "UC-03 CRC placeholder not found"
print(f"UC-03 CRC placeholder found: {repr(gettxt(uc03_crc_placeholder)[:50])}")

crc_elems = []
crc_elems.append(make_para2(
    "Thẻ CRC (Class–Responsibility–Collaborator) được lập cho 5 lớp lĩnh vực "
    "trung tâm của UC-03. Các lớp TaiXe và DonHang được lập thẻ CRC ở UC-02 và UC-01 "
    "tương ứng vì chúng xuất hiện lần đầu ở đó. Thuộc tính chỉ ghi tên — không có "
    "kiểu dữ liệu (quy tắc phân tích I.12)."
))
crc_elems.append(make_para2())

crc_data = [
    ("ChuyenDi", "CRC-UC03-01", "Lĩnh vực (Domain)",
     "Đại diện cho một ca làm việc của tài xế, bao gồm một hoặc nhiều lộ trình giao hàng trong một ngày.",
     "maChuyenDi, thoiGianBatDau, thoiGianKetThuc",
     "1. Ghi nhận thời điểm bắt đầu và kết thúc ca làm việc.\n"
     "2. Liên kết tài xế với tập hợp các lộ trình trong ca.\n"
     "3. Cung cấp ngữ cảnh ca làm việc để tổng hợp hiệu suất.",
     "TaiXe, LoTrinh"),
    ("LoTrinh", "CRC-UC03-02", "Lĩnh vực (Domain)",
     "Tuyến đường từ kho xuất phát đến điểm giao hàng cho một đơn cụ thể; "
     "lưu trữ chuỗi tọa độ GPS toàn hành trình.",
     "maLoTrinh, diemXuatPhat, diemDen, khoangCach, /thoiGianDuKien, trangThai",
     "1. Lưu thông tin tuyến đường (điểm đầu, điểm cuối, khoảng cách).\n"
     "2. Tập hợp các bản ghi ViTriGPS trong suốt hành trình.\n"
     "3. Theo dõi trạng thái: Đang giao → Hoàn tất.\n"
     "4. Cung cấp dữ liệu ETA cho ứng dụng tài xế.",
     "ChuyenDi, TaiXe, DonHang, ViTriGPS"),
    ("ViTriGPS", "CRC-UC03-03", "Lĩnh vực (Domain)",
     "Một bản ghi tọa độ GPS tại một thời điểm cụ thể trong hành trình; được ghi nhận mỗi 30 giây.",
     "maViTri, viDo, kinhDo, thoiGian",
     "1. Lưu tọa độ địa lý (vĩ độ, kinh độ) tại thời điểm ghi nhận.\n"
     "2. Cung cấp dữ liệu cho tính năng theo dõi vị trí thời gian thực.\n"
     "3. Tạo vết hành trình (breadcrumb trail) để kiểm tra sau giao.",
     "LoTrinh"),
    ("KienHang", "CRC-UC03-04", "Lĩnh vực (Domain)",
     "Đơn vị vật lý được vận chuyển; được xác nhận bằng mã QR khi tài xế lấy hàng tại kho.",
     "maKienHang, maQR, khoiLuong, trangThai",
     "1. Xác minh danh tính kiện hàng qua quét mã QR.\n"
     "2. Theo dõi trạng thái vật lý: Đang vận chuyển → Đã giao.\n"
     "3. Cung cấp thông tin khoiLuong để tính phí vận chuyển (UC-05).",
     "DonHang"),
    ("BangChungGiaoHang", "CRC-UC03-05", "Lĩnh vực (Domain)",
     "Bằng chứng xác thực việc giao hàng thành công; lưu ảnh chụp biên nhận, "
     "tên người nhận và tọa độ GPS tại điểm giao.",
     "maBangChung, tenNguoiNhan, hinhAnh, toaDo, thoiGianGiao",
     "1. Lưu trữ ảnh biên nhận giao hàng (hinhAnh).\n"
     "2. Ghi nhận tên người nhận thực tế và vị trí giao hàng.\n"
     "3. Cung cấp bằng chứng cho quá trình xác nhận giao hàng ở UC-04.",
     "DonHang, XacNhanGiaoHang"),
]

for args in crc_data:
    t, s = make_crc_table(*args)
    crc_elems.extend([t, s])

insert_before(crc_elems, uc03_crc_placeholder)
uc03_crc_placeholder.getparent().remove(uc03_crc_placeholder)
print(f"Inserted {len(crc_elems)} CRC elements into UC-03 4.3.5")

doc.save(DOC_PATH)
print("Done.")
