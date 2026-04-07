"""
Inserts a new Section 3 — "Phân tích danh từ và xác định lớp lĩnh vực"
into Ch4_MoHinhHoaCauTruc.docx, immediately before the '4.1. UC-01' heading.

Content sourced from docs/ch4-draft.md sections 4.2.1 – 4.2.4
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import copy

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"

doc = Document(DOC_PATH)
body = doc.element.body

# ── helpers ───────────────────────────────────────────────────────────────────

def ft(para):
    return "".join(r.text for r in para.runs) or para.text

def detach(para_or_tbl):
    """Remove an element from its parent and return the element."""
    e = para_or_tbl._element if hasattr(para_or_tbl, "_element") else para_or_tbl
    e.getparent().remove(e)
    return e

def make_para(style="Normal", text="", bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return detach(p)

def make_heading1(text):
    return make_para("Heading 1", text)

def make_heading2(text):
    return make_para("Heading 2", text)

def make_heading3(text):
    return make_para("Heading 3", text)

def make_bold_intro(label, rest):
    """Bold label + normal rest in one paragraph."""
    p = doc.add_paragraph(style="Normal")
    r1 = p.add_run(label)
    r1.bold = True
    if rest:
        p.add_run(rest)
    return detach(p)

def make_table(headers, rows, col_widths=None):
    """Returns detached table element + trailing spacer paragraph element."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Use first available table style
    try:
        tbl.style = "TableNormal"
    except Exception:
        pass

    # Header row — bold
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)

    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def make_bullet(text):
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except Exception:
        p = doc.add_paragraph("– " + text)
    return detach(p)

def insert_before_ref(elems, ref_elem):
    """Insert list of elements before ref_elem in body."""
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)

# ── find insertion point: before '4.1. UC-01' ─────────────────────────────────

insert_ref = None
for p in doc.paragraphs:
    if p.style.name == "Heading 2" and ft(p).startswith("4.1."):
        insert_ref = p._element
        break

if insert_ref is None:
    raise RuntimeError("Could not find '4.1.' heading as insertion anchor.")

# ── build all new elements ─────────────────────────────────────────────────────

elems = []

# ═══════════════════════════════════════════════════════════════════
# Section 3 — Phân tích danh từ và xác định lớp lĩnh vực
# ═══════════════════════════════════════════════════════════════════
elems.append(make_heading1("3. Phân tích danh từ và xác định lớp lĩnh vực"))
elems.append(make_para("Normal",
    "Trước khi xây dựng sơ đồ lớp cho từng ca sử dụng, nhóm thực hiện "
    "phân tích danh từ trên toàn bộ đặc tả 5 UC để nhận diện các khái niệm "
    "nghiệp vụ cốt lõi. Kết quả của bước này là danh sách 15 lớp lĩnh vực, "
    "tập đối tượng mẫu và ma trận liên kết — nền tảng thống nhất cho tất cả "
    "các sơ đồ trong Mục 4."
))

# ─── 3.1 Kỹ thuật danh từ ─────────────────────────────────────────
elems.append(make_heading2("3.1. Kỹ thuật danh từ (Noun Technique)"))
elems.append(make_para("Normal",
    "Nhóm đọc toàn bộ đặc tả luồng sự kiện chính và phụ của 5 UC, gạch chân "
    "mọi danh từ — kể cả những thứ có vẻ hiển nhiên. Danh từ có thể là tên "
    "lớp, tên thuộc tính, hay nằm ngoài phạm vi hệ thống. Nguồn dữ liệu bao "
    "gồm: (1) đặc tả 5 UC trong Chương III; (2) phiếu xuất kho, hóa đơn COD "
    "và bảng điểm tài xế của LogiFast; (3) ghi chép phỏng vấn điều phối viên "
    "và tài xế."
))
elems.append(make_para("Normal", "Bảng 3.1 — Danh sách danh từ thô trích từ 5 UC LogiFast:", bold=True))

noun_headers = ["STT", "Danh từ thô", "Xuất hiện trong"]
noun_rows = [
    ["1",  "Khách hàng",               "UC-01, UC-04"],
    ["2",  "Đơn hàng",                 "UC-01, UC-02, UC-03, UC-04, UC-05"],
    ["3",  "Giỏ hàng",                 "UC-01"],
    ["4",  "Sản phẩm / mặt hàng",      "UC-01"],
    ["5",  "Kiện hàng",                "UC-01, UC-03"],
    ["6",  "Địa chỉ giao hàng",        "UC-01"],
    ["7",  "Phương thức thanh toán",   "UC-01, UC-05"],
    ["8",  "Mã đơn hàng",              "UC-01, UC-04"],
    ["9",  "Trang xác nhận",           "UC-01"],
    ["10", "Thời gian giao dự kiến",   "UC-01, UC-03"],
    ["11", "Tài xế",                   "UC-02, UC-03"],
    ["12", "Phiếu phân công",          "UC-02"],
    ["13", "Vùng địa chỉ",             "UC-02"],
    ["14", "Điểm hiệu suất",           "UC-02"],
    ["15", "Bán kính",                 "UC-02"],
    ["16", "Kho hàng",                 "UC-02, UC-03"],
    ["17", "Lộ trình",                 "UC-03"],
    ["18", "Vị trí GPS",               "UC-03"],
    ["19", "Mã QR",                    "UC-03"],
    ["20", "Chuyến đi",                "UC-03"],
    ["21", "Bằng chứng giao hàng",     "UC-03, UC-04"],
    ["22", "Ảnh chụp",                 "UC-03, UC-04"],
    ["23", "Tên người nhận",           "UC-03, UC-04"],
    ["24", "Shipper",                  "UC-04"],
    ["25", "Xác nhận giao hàng",       "UC-04"],
    ["26", "Chữ ký điện tử / OTP",    "UC-04"],
    ["27", "Thông báo",                "UC-04"],
    ["28", "Giao dịch",                "UC-05"],
    ["29", "Hóa đơn",                  "UC-05"],
    ["30", "Ví điện tử",               "UC-05"],
    ["31", "Sổ cái",                   "UC-05"],
    ["32", "Phí dịch vụ",              "UC-05"],
    ["33", "Hoa hồng tài xế",          "UC-05"],
    ["34", "Hệ thống",                 "Tất cả UC"],
    ["35", "Ứng dụng di động",         "Tất cả UC"],
    ["36", "Quản lý / Điều phối viên", "UC-02"],
]
t, s = make_table(noun_headers, noun_rows)
elems.extend([t, s])

# ─── 3.2 Phân loại danh từ ────────────────────────────────────────
elems.append(make_heading2("3.2. Phân loại danh từ"))
elems.append(make_para("Normal",
    "Với mỗi danh từ, nhóm áp dụng bộ câu hỏi quyết định để phân loại thành "
    "lớp, thuộc tính, thuộc tính suy diễn, đồng nghĩa, hoặc ngoài phạm vi:"
))

q_headers = ["Câu hỏi quyết định", "Nếu \"Có\"", "Nếu \"Không\""]
q_rows = [
    ["Danh từ có nằm trong phạm vi hệ thống LogiFast?",
     "Giữ lại, xét tiếp", "Loại bỏ (ngoài phạm vi)"],
    ["Hệ thống cần ghi nhớ nhiều hơn 1 đối tượng loại này?",
     "Lớp ứng viên", "Có thể là hằng số / cấu hình"],
    ["Nó là thành phần của thứ khác đã xác định?",
     "Thuộc tính", "Giữ như lớp độc lập"],
    ["Nó là đồng nghĩa với thứ khác đã xác định?",
     "Loại bỏ / gộp", "Tiếp tục xét"],
    ["Nó chỉ là đầu ra được tính từ dữ liệu khác?",
     "Thuộc tính suy diễn (/tên)", "Giữ như thuộc tính thường"],
]
t, s = make_table(q_headers, q_rows)
elems.extend([t, s])

elems.append(make_para("Normal", "Bảng 3.2 — Kết quả phân loại toàn bộ danh từ thô LogiFast:", bold=True))

cls_headers = ["Danh từ thô", "Phân loại", "Ghi chú / Kết quả"]
cls_rows = [
    ["Khách hàng",             "Lớp",                   "KhachHang — khách B2C và đối tác B2B qua API"],
    ["Đơn hàng",               "Lớp",                   "DonHang — trung tâm toàn bộ quy trình"],
    ["Giỏ hàng",               "Ngoài phạm vi",          "LogiFast không quản lý giỏ hàng — thuộc hệ thống đối tác"],
    ["Sản phẩm / mặt hàng",    "Lớp",                   "SanPham — cần ghi nhớ để tính phí"],
    ["Kiện hàng",              "Lớp",                   "KienHang — đơn vị vật lý được vận chuyển"],
    ["Địa chỉ giao hàng",      "Thuộc tính",             "diaChiGiao trong DonHang (đơn trị)"],
    ["Phương thức thanh toán", "Lớp",                   "PhuongThucThanhToan — COD, prepaid, B2B credit"],
    ["Mã đơn hàng",            "Thuộc tính",             "maDonHang trong DonHang — khóa định danh"],
    ["Trang xác nhận",         "Ngoài phạm vi",          "Thành phần UI, không phải lớp lĩnh vực"],
    ["Thời gian giao dự kiến", "Thuộc tính suy diễn",   "/thoiGianDuKien trong LoTrinh"],
    ["Tài xế",                 "Lớp",                   "TaiXe — agent vận chuyển"],
    ["Phiếu phân công",        "Lớp",                   "PhieuPhanCong — lớp liên kết DonHang <-> TaiXe"],
    ["Vùng địa chỉ",           "Lớp",                   "VungDiaChi — dùng trong thuật toán phân công"],
    ["Điểm hiệu suất",         "Lớp",                   "DiemHieuSuat — gắn với TaiXe, cần lịch sử"],
    ["Bán kính",               "Thuộc tính",             "banKinhTimKiem trong VungDiaChi"],
    ["Kho hàng",               "Thuộc tính",             "diemXuatPhat trong LoTrinh (không cần lớp riêng)"],
    ["Lộ trình",               "Lớp",                   "LoTrinh — UC-03 phức tạp nhất"],
    ["Vị trí GPS",             "Lớp",                   "ViTriGPS — cần lưu lịch sử chuỗi tọa độ"],
    ["Mã QR",                  "Thuộc tính",             "maQR trong KienHang"],
    ["Chuyến đi",              "Lớp",                   "ChuyenDi — bao gồm nhiều đơn trong 1 ca làm việc"],
    ["Bằng chứng giao hàng",   "Lớp",                   "BangChungGiaoHang — dùng chung UC-03 và UC-04"],
    ["Ảnh chụp",               "Thuộc tính",             "hinhAnh trong BangChungGiaoHang"],
    ["Tên người nhận",         "Thuộc tính",             "tenNguoiNhan trong BangChungGiaoHang"],
    ["Shipper",                "Đồng nghĩa",             "Tương đương TaiXe trong phạm vi 5 UC — không tạo lớp riêng"],
    ["Xác nhận giao hàng",     "Lớp",                   "XacNhanGiaoHang — hành động có dữ liệu cần ghi nhớ"],
    ["OTP / Chữ ký điện tử",  "Thuộc tính",             "phuongThucXacNhan trong XacNhanGiaoHang"],
    ["Thông báo",              "Lớp",                   "ThongBao — UC-04 gửi đến nhiều bên"],
    ["Giao dịch",              "Lớp",                   "GiaoDich — UC-05"],
    ["Hóa đơn",                "Lớp",                   "HoaDon — tín dụng B2B, cần ghi nhớ riêng"],
    ["Ví điện tử",             "Lớp",                   "ViDienTu — thuộc TaiXe"],
    ["Sổ cái",                 "Lớp",                   "SoCai — ghi nhận doanh thu tổng"],
    ["Phí dịch vụ",            "Lớp",                   "PhiDichVu — loại phí có cấu trúc riêng"],
    ["Hoa hồng tài xế",        "Thuộc tính suy diễn",   "/hoaHong trong GiaoDich"],
    ["Hệ thống",               "Loại bỏ",               "Quá chung, không phải lớp lĩnh vực"],
    ["Ứng dụng di động",       "Loại bỏ",               "Thành phần UI"],
    ["Quản lý / Điều phối viên","Loại bỏ",              "Actor — không ghi nhớ trong mô hình lĩnh vực"],
]
t, s = make_table(cls_headers, cls_rows)
elems.extend([t, s])

elems.append(make_para("Normal",
    "Kết quả: 15 lớp lĩnh vực được xác định (KhachHang, DonHang, SanPham, "
    "KienHang, PhuongThucThanhToan, TaiXe, PhieuPhanCong, VungDiaChi, DiemHieuSuat, "
    "LoTrinh, ViTriGPS, ChuyenDi, BangChungGiaoHang, XacNhanGiaoHang, ThongBao, "
    "GiaoDich, HoaDon, ViDienTu, SoCai, PhiDichVu). "
    "Phân phối theo UC được trình bày chi tiết trong Mục 4."
))

# ─── 3.3 Xây dựng các đối tượng ──────────────────────────────────
elems.append(make_heading2("3.3. Xây dựng các đối tượng (Object Instantiation)"))
elems.append(make_para("Normal",
    "Với mỗi UC, nhóm xây dựng một kịch bản cụ thể có giá trị thuộc tính thực tế, "
    "sau đó tạo các đối tượng (instance) cho mỗi lớp tham gia. "
    "Một đối tượng UML được ký hiệu tenDoiTuong : TenLop (gạch chân). "
    "Các quy tắc áp dụng: (1) mỗi đối tượng phải có tên instance và tên lớp; "
    "(2) giá trị thuộc tính phải cụ thể và thực tế; "
    "(3) đối tượng xuất hiện trong nhiều UC phải dùng cùng giá trị ID; "
    "(4) cơ số trong sơ đồ lớp phải được thỏa mãn trong sơ đồ đối tượng."
))
elems.append(make_para("Normal",
    "Bảng 3.3 — Instance catalog cho kịch bản happy path LogiFast (26/03/2026):",
    bold=True
))

inst_headers = ["Tên đối tượng", "Lớp", "UC xuất hiện", "ID mẫu"]
inst_rows = [
    ["kh1",      "KhachHang",          "UC-01, UC-04",         "KH-055"],
    ["dh1",      "DonHang",            "UC-01 đến UC-05",      "ORD-2026-001"],
    ["sp1",      "SanPham",            "UC-01",                "SP-0099"],
    ["ctd1",     "ChiTietDonHang",     "UC-01",                "CTD-001-01"],
    ["ptt1",     "PhuongThucThanhToan","UC-01, UC-05",         "PTT-COD"],
    ["tx1",      "TaiXe",              "UC-02, UC-03, UC-04",  "TX-001"],
    ["ppc1",     "PhieuPhanCong",      "UC-02",                "PPC-2026-0326-001"],
    ["vda1",     "VungDiaChi",         "UC-02",                "VDA-Q5-BCHB"],
    ["dhs1",     "DiemHieuSuat",       "UC-02",                "DHS-TX-001"],
    ["kh_pkg1",  "KienHang",           "UC-01, UC-03",         "KH-PKG-001"],
    ["lt1",      "LoTrinh",            "UC-03",                "LT-2026-001"],
    ["gps1",     "ViTriGPS",           "UC-03",                "GPS-TX001-1006"],
    ["cd1",      "ChuyenDi",           "UC-03",                "CD-2026-0326-T1"],
    ["bc1",      "BangChungGiaoHang",  "UC-03, UC-04",         "BC-2026-001"],
    ["xn1",      "XacNhanGiaoHang",    "UC-04",                "XN-2026-001"],
    ["tb1",      "ThongBao",           "UC-04",                "TB-2026-001"],
    ["gd1",      "GiaoDich",           "UC-05",                "GD-2026-001"],
    ["hd1",      "HoaDon",             "UC-05",                "HD-2026-001"],
    ["vdt1",     "ViDienTu",           "UC-05",                "VDT-TX001"],
    ["sc1",      "SoCai",              "UC-05",                "SC-2026-Q1"],
    ["phi1",     "PhiDichVu",          "UC-05",                "PHI-2026-001"],
]
t, s = make_table(inst_headers, inst_rows)
elems.extend([t, s])

# ─── 3.4 Liên kết đối tượng ──────────────────────────────────────
elems.append(make_heading2("3.4. Liên kết đối tượng (Object Linking)"))
elems.append(make_para("Normal",
    "Từ kịch bản thực tế, nhóm xác định liên kết giữa các đối tượng bằng cách "
    "đọc động từ kết nối danh từ trong luồng sự kiện của từng UC: "
    "\"tài xế nhận đơn hàng\", \"đơn hàng chứa kiện hàng\", v.v. "
    "Mỗi liên kết đối tượng (link) phải tương ứng với một quan hệ (association) "
    "trong sơ đồ lớp lĩnh vực. Nếu phát hiện liên kết không có quan hệ tương "
    "ứng, sơ đồ lớp cần được bổ sung. Cơ số của quan hệ phải được thỏa mãn "
    "trong ví dụ đối tượng tương ứng."
))
elems.append(make_para("Normal",
    "Bảng 3.4 — Ma trận liên kết đối tượng — happy path toàn quy trình:",
    bold=True
))

link_headers = ["Từ đối tượng", "Đến đối tượng", "Tên liên kết", "Nguồn (UC)"]
link_rows = [
    ["kh1",      "dh1",      "đặt",            "UC-01"],
    ["dh1",      "ctd1",     "gồm",            "UC-01"],
    ["ctd1",     "sp1",      "tham chiếu",     "UC-01"],
    ["dh1",      "ptt1",     "sử dụng",        "UC-01"],
    ["dh1",      "kh_pkg1",  "chứa",           "UC-01, UC-03"],
    ["ppc1",     "dh1",      "phân công",      "UC-02"],
    ["ppc1",     "tx1",      "gán",            "UC-02"],
    ["tx1",      "vda1",     "thuộc vùng",     "UC-02"],
    ["tx1",      "dhs1",     "có",             "UC-02"],
    ["cd1",      "lt1",      "bao gồm",        "UC-03"],
    ["tx1",      "lt1",      "thực hiện",      "UC-03"],
    ["lt1",      "dh1",      "giao",           "UC-03"],
    ["lt1",      "gps1",     "ghi nhận",       "UC-03"],
    ["dh1",      "bc1",      "có bằng chứng",  "UC-03, UC-04"],
    ["tx1",      "xn1",      "lập",            "UC-04"],
    ["bc1",      "xn1",      "đính kèm",       "UC-04"],
    ["xn1",      "tb1",      "kích hoạt",      "UC-04"],
    ["dh1",      "gd1",      "kết toán",       "UC-05"],
    ["gd1",      "hd1",      "tạo",            "UC-05"],
    ["gd1",      "phi1",     "áp dụng",        "UC-05"],
    ["tx1",      "vdt1",     "sở hữu",         "UC-05"],
    ["gd1",      "vdt1",     "cộng vào",       "UC-05"],
    ["gd1",      "sc1",      "ghi vào",        "UC-05"],
]
t, s = make_table(link_headers, link_rows)
elems.extend([t, s])

elems.append(make_para("Normal",
    "Toàn bộ 23 liên kết này đều có quan hệ lớp tương ứng trong các sơ đồ lớp "
    "lĩnh vực theo UC (Mục 4). Không phát hiện liên kết mồ côi trong kịch bản "
    "happy path."
))

# Blank spacer before section 4
elems.append(make_para("Normal", ""))

# ── insert all elements before '4.1.' ─────────────────────────────────────────
insert_before_ref(elems, insert_ref)
print(f"Inserted {len(elems)} elements before '4.1. UC-01'")

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Done. Saved to", DOC_PATH)
