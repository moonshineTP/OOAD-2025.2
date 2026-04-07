"""
Replaces the 4.3.1 section content (between Heading3 '4.3.1.' and Heading3 '4.3.2.')
with a comprehensive UC-03 narrative covering:
  1. Happy scenario (7 timestamped steps)
  2. Noun analysis table (UC-03 specific)
  3. Object instantiation table (7 objects, full attribute values)
  4. Object linking table (6 links)
"""

from docx import Document
from docx.oxml import OxmlElement
import copy

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document(DOC_PATH)
body = doc.element.body

# ── helpers ───────────────────────────────────────────────────────────────────

def get_text(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{NS}}}t"))

def detach(para_or_tbl):
    e = para_or_tbl._element if hasattr(para_or_tbl, "_element") else para_or_tbl
    e.getparent().remove(e)
    return e

def make_para(text="", style="Normal", bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return detach(p)

def make_mixed(segments):
    """segments = list of (text, bold, italic)"""
    p = doc.add_paragraph(style="Normal")
    for text, bold, italic in segments:
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
    return detach(p)

def make_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = "TableNormal"
    except Exception:
        pass
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
    tbl_elem = detach(tbl)
    spacer = detach(doc.add_paragraph())
    return tbl_elem, spacer

def make_bullet(text):
    try:
        p = doc.add_paragraph(style="List Bullet")
    except Exception:
        p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return detach(p)

def insert_before(elems, ref_elem):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for i, e in enumerate(elems):
        parent.insert(idx + i, e)

# ── find bounds: body children between 4.3.1 heading and 4.3.2 heading ────────

body_children = list(body)
start_idx = None  # index of 4.3.1 heading (keep it)
end_idx = None    # index of 4.3.2 heading (keep it — insert before)
first_content_idx = None  # first child AFTER heading to remove

for i, child in enumerate(body_children):
    tag = child.tag.split("}")[-1]
    text = get_text(child)
    if tag == "p":
        style_elem = child.find(f".//{{{NS}}}pStyle")
        sval = style_elem.get(f"{{{NS}}}val") if style_elem is not None else ""
        if "4.3.1" in text and "Heading" in sval:
            start_idx = i
            first_content_idx = i + 1
        if "4.3.2" in text and "Heading" in sval and start_idx is not None:
            end_idx = i
            break

print(f"start={start_idx}, first_content={first_content_idx}, end={end_idx}")

# Remove all children between first_content_idx and end_idx
to_remove = body_children[first_content_idx:end_idx]
for elem in to_remove:
    body.remove(elem)
print(f"Removed {len(to_remove)} elements")

# Reference for insertion = 4.3.2 heading (now shifted)
ref_elem = list(body)[start_idx + 1]  # first child after 4.3.1 heading

# ── build new content ─────────────────────────────────────────────────────────

new_elems = []

# ── A. Tên kịch bản ───────────────────────────────────────────────────────────
new_elems.append(make_mixed([
    ("Tên kịch bản: ", True, False),
    ("Tài xế TX-001 lấy hàng tại kho và giao thành công đến khách hàng KH-055", False, False),
]))

# ── B. Diễn giải (7-step timestamped narrative) ───────────────────────────────
new_elems.append(make_para())  # blank

new_elems.append(make_mixed([("Diễn giải kịch bản:", True, False)]))

steps = [
    ("09:28", (
        "TX-001 (Trần Minh Khoa) mở ứng dụng LogiFast Driver và xác nhận bắt đầu ca làm việc. "
        "Hệ thống tạo ChuyenDi CD-2026-0326-T1 và giao nhiệm vụ: lấy đơn ORD-2026-001 tại kho "
        "KHO-SGN-01 (Kho Bình Chánh, Đường số 12, KCN Vĩnh Lộc)."
    )),
    ("09:30", (
        "TX-001 đến kho. Nhân viên kho xuất kiện hàng KH-PKG-001. TX-001 quét mã QR "
        "\"QR-ORD-2026-001\" bằng ứng dụng. Hệ thống xác nhận đúng đơn: Samsung Galaxy A56, "
        "0,35 kg, địa chỉ giao \"123 Nguyễn Văn Cừ, Q.5\". "
        "KienHang.trangThai → \"Đang vận chuyển\"."
    )),
    ("09:31", (
        "Hệ thống tính lộ trình tối ưu LT-2026-001: khoảng cách 9,2 km, "
        "thời gian dự kiến 35 phút (ETA 10:06). TX-001 nhấn \"Bắt đầu giao hàng\". "
        "DonHang.trangThai → \"Đang giao\". ChuyenDi.thoiGianBatDau = \"09:30\"."
    )),
    ("09:31 → 10:04", (
        "Tài xế di chuyển theo lộ trình được chỉ định. Hệ thống ghi nhận ViTriGPS mỗi 30 giây "
        "(khoảng 40 điểm GPS trong suốt hành trình, tất cả liên kết vào LoTrinh lt1). "
        "Điểm GPS cuối được ghi: gps1 — tọa độ (10,7525; 106,6624) lúc 10:06:00."
    )),
    ("10:04", (
        "TX-001 cách địa chỉ giao khoảng 200m. Hệ thống tự động gửi thông báo "
        "\"Shipper đang đến, ETA 2 phút\" đến KH-055 (0901234567) qua ứng dụng và SMS."
    )),
    ("10:06", (
        "TX-001 gặp chị Nguyễn Thị Lan — người nhận đơn — tại cửa \"123 Nguyễn Văn Cừ, Q.5\". "
        "Bàn giao kiện hàng thành công. Trên ứng dụng, TX-001 chụp ảnh biên nhận và nhập tên người nhận. "
        "Hệ thống tạo BangChungGiaoHang BC-2026-001: lưu ảnh proof_ORD-2026-001.jpg, "
        "tọa độ GPS (10,7525; 106,6624) và tên người nhận \"Nguyễn Thị Lan\"."
    )),
    ("10:06 (hoàn tất)", (
        "Hệ thống cập nhật: DonHang.trangThai → \"Giao thành công\"; "
        "KienHang.trangThai → \"Đã giao\"; LoTrinh.trangThai → \"Hoàn tất\". "
        "Sự kiện \"DeliveryCompleted\" được phát ra để kích hoạt UC-04 (Xác nhận giao hàng)."
    )),
]

# Build timestep paragraphs
for timestamp, text in steps:
    new_elems.append(make_mixed([
        (f"{timestamp} — ", True, False),
        (text, False, False),
    ]))

new_elems.append(make_para())  # blank

# ── C. Phân tích danh từ UC-03 ────────────────────────────────────────────────
new_elems.append(make_mixed([("Phân tích danh từ UC-03:", True, False)]))
new_elems.append(make_para(
    "Áp dụng kỹ thuật danh từ vào đặc tả UC-03, nhóm xác định các danh từ xuất hiện "
    "trực tiếp trong luồng sự kiện vận chuyển và phân loại như sau:"
))

noun_h = ["Danh từ thô (UC-03)", "Phân loại", "Kết quả"]
noun_r = [
    ["Tài xế",                "Lớp",                 "TaiXe — agent thực hiện vận chuyển"],
    ["Đơn hàng",              "Lớp",                 "DonHang — kế thừa từ UC-01/UC-02"],
    ["Kiện hàng",             "Lớp",                 "KienHang — đơn vị vật lý được vận chuyển"],
    ["Mã QR",                 "Thuộc tính",           "maQR trong KienHang"],
    ["Lộ trình",              "Lớp",                 "LoTrinh — tuyến đường từ kho đến địa chỉ giao"],
    ["Thời gian giao dự kiến","Thuộc tính suy diễn", "/thoiGianDuKien trong LoTrinh (tính từ khoangCach)"],
    ["Kho hàng",              "Thuộc tính",           "diemXuatPhat trong LoTrinh (không cần lớp riêng)"],
    ["Vị trí GPS",            "Lớp",                 "ViTriGPS — cần ghi nhớ chuỗi tọa độ theo thời gian"],
    ["Chuyến đi",             "Lớp",                 "ChuyenDi — một ca làm việc bao gồm nhiều lộ trình"],
    ["Bằng chứng giao hàng",  "Lớp",                 "BangChungGiaoHang — lưu ảnh, GPS, tên người nhận"],
    ["Ảnh chụp",              "Thuộc tính",           "hinhAnh trong BangChungGiaoHang"],
    ["Tên người nhận",        "Thuộc tính",           "tenNguoiNhan trong BangChungGiaoHang"],
    ["Thông báo đến khách",   "Lớp (dùng lại)",      "ThongBao — tái sử dụng từ UC-04"],
    ["Ứng dụng LogiFast Driver","Loại bỏ",           "Thành phần UI — không phải lớp lĩnh vực"],
]
t, s = make_table(noun_h, noun_r)
new_elems.extend([t, s])

new_elems.append(make_para(
    "UC-03 xác định 5 lớp lĩnh vực mới (LoTrinh, ViTriGPS, ChuyenDi, BangChungGiaoHang và "
    "tái sử dụng KienHang lần đầu ở mức vận chuyển) không xuất hiện trong UC-01 hay UC-02, "
    "đảm bảo tiêu chí I.11 về sự khác biệt có ý nghĩa giữa các sơ đồ lớp."
))
new_elems.append(make_para())

# ── D. Xây dựng đối tượng UC-03 ──────────────────────────────────────────────
new_elems.append(make_mixed([("Bảng đối tượng UC-03:", True, False)]))

obj_h = ["Lớp", "Tên đối tượng", "Thuộc tính chính (giá trị kịch bản)"]
obj_r = [
    ["TaiXe",
     "tx1 : TaiXe",
     "id=TX-001, ten=\"Trần Minh Khoa\", soDienThoai=\"0912345678\", "
     "trangThai=\"Đã đến điểm giao\", viTriHienTai=\"10.7525, 106.6624\""],
    ["DonHang",
     "dh1 : DonHang",
     "id=ORD-2026-001, trangThai=\"Giao thành công\", "
     "diaChiGiao=\"123 Nguyễn Văn Cừ, Q.5\", tongGiaTri=8.500.000 VND"],
    ["KienHang",
     "kh_pkg1 : KienHang",
     "id=KH-PKG-001, maQR=\"QR-ORD-2026-001\", khoiLuong=0.35 kg, trangThai=\"Đã giao\""],
    ["LoTrinh",
     "lt1 : LoTrinh",
     "id=LT-2026-001, diemXuatPhat=\"KHO-SGN-01\", "
     "diemDen=\"123 Nguyễn Văn Cừ, Q.5\", khoangCach=9.2 km, "
     "/thoiGianDuKien=35 phút, trangThai=\"Hoàn tất\""],
    ["ViTriGPS",
     "gps1 : ViTriGPS",
     "id=GPS-TX001-1006, viDo=10.7525, kinhDo=106.6624, "
     "thoiGian=\"2026-03-26T10:06:00\" (điểm cuối hành trình)"],
    ["ChuyenDi",
     "cd1 : ChuyenDi",
     "id=CD-2026-0326-T1, thoiGianBatDau=\"2026-03-26T09:28:00\", "
     "thoiGianKetThuc=\"2026-03-26T10:10:00\""],
    ["BangChungGiaoHang",
     "bc1 : BangChungGiaoHang",
     "id=BC-2026-001, tenNguoiNhan=\"Nguyễn Thị Lan\", "
     "hinhAnh=\"proof_ORD-2026-001.jpg\", "
     "toaDo=\"10.7525, 106.6624\", thoiGianGiao=\"2026-03-26T10:06:00\""],
]
t, s = make_table(obj_h, obj_r)
new_elems.extend([t, s])

# ── E. Liên kết đối tượng ────────────────────────────────────────────────────
new_elems.append(make_mixed([("Liên kết đối tượng UC-03:", True, False)]))
new_elems.append(make_para(
    "Từ luồng sự kiện vận chuyển, các liên kết sau được xác định bằng cách đọc "
    "động từ kết nối danh từ trong từng bước:"
))

link_h = ["Từ", "Đến", "Tên liên kết", "Bước UC tương ứng"]
link_r = [
    ["cd1 : ChuyenDi",            "lt1 : LoTrinh",             "bao gồm",       "Bước 1 — tạo ChuyenDi + LoTrinh"],
    ["tx1 : TaiXe",               "lt1 : LoTrinh",             "thực hiện",     "Bước 3 — tài xế xuất phát theo lộ trình"],
    ["lt1 : LoTrinh",             "dh1 : DonHang",             "giao",          "Bước 3–6 — lộ trình phục vụ giao đơn"],
    ["lt1 : LoTrinh",             "gps1 : ViTriGPS",           "ghi nhận",      "Bước 4 — GPS cập nhật liên tục"],
    ["dh1 : DonHang",             "kh_pkg1 : KienHang",        "chứa",          "Bước 2 — quét QR xác nhận kiện hàng"],
    ["dh1 : DonHang",             "bc1 : BangChungGiaoHang",   "có bằng chứng","Bước 6 — tạo BC sau khi giao xong"],
]
t, s = make_table(link_h, link_r)
new_elems.extend([t, s])

new_elems.append(make_para(
    "Tất cả 6 liên kết trên phải có quan hệ tương ứng trong sơ đồ lớp lĩnh vực UC-03 "
    "(Mục 4.3.3). Cơ số được thỏa mãn: cd1 → lt1 (1 ChuyenDi có 1 LoTrinh trong kịch bản này); "
    "lt1 → gps1 đại diện cho 1 trong ~40 điểm GPS được ghi nhận (quan hệ 1..*); "
    "dh1 → kh_pkg1 thỏa mãn cơ số 1..* (đơn hàng phải có ít nhất 1 kiện hàng)."
))

# ── insert ────────────────────────────────────────────────────────────────────
insert_before(new_elems, ref_elem)
print(f"Inserted {len(new_elems)} elements")

doc.save(DOC_PATH)
print("Done.")
