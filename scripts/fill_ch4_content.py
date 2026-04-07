"""
Fills in the natural-language scenario content of Ch4_MoHinhHoaCauTruc.docx.

For each UC section (4.1–4.5), replaces the placeholder paragraph
  "[Mô tả kịch bản cụ thể và bảng đối tượng]"
with:
  • Tên kịch bản (bold label + text)
  • Diễn giải (narrative paragraph)
  • "Bảng đối tượng" heading paragraph
  • A Word table with the object data
  • "Quan hệ đối tượng" heading + bullet list

Diagram placeholders [Chèn sơ đồ ...] are left untouched.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

DOC_PATH = r"Project/Ch4_MoHinhHoaCauTruc.docx"

# ─── Content data per UC ────────────────────────────────────────────────────

UC_SCENARIOS = {
    "4.1": {
        "title": "Kịch bản: Khách hàng B2C đặt đơn giao 1 kiện hàng",
        "narrative": (
            "Vào 09:15 ngày 26/03/2026, chị Nguyễn Thị Lan (KH-055) đăng nhập ứng dụng "
            "LogiFast Mobile, xác nhận giỏ hàng gồm 1 sản phẩm (Điện thoại Samsung A56, "
            "SP-0099), nhập địa chỉ giao \"123 Nguyễn Văn Cừ, Q.5, TP.HCM\" và chọn hình "
            "thức thanh toán COD. Hệ thống kiểm tra tính hợp lệ của đơn hàng, tạo mã đơn "
            "duy nhất ORD-2026-001 với trạng thái Sẵn sàng giao, rồi phát sự kiện "
            "\"OrderCreated\" để kích hoạt quy trình phân công tự động."
        ),
        "table_headers": ["Lớp", "Tên đối tượng", "Giá trị thuộc tính chính"],
        "table_rows": [
            ["KhachHang", "kh1 : KhachHang",
             "id=KH-055, ten=\"Nguyễn Thị Lan\", soDienThoai=\"0901234567\", diaChi=\"45 Lê Lợi, Q.1\", loaiKhach=\"B2C\""],
            ["DonHang", "dh1 : DonHang",
             "id=ORD-2026-001, trangThai=\"Sẵn sàng giao\", diaChiGiao=\"123 Nguyễn Văn Cừ, Q.5\", tongGiaTri=8.500.000 VND, hinhThucThanhToan=\"COD\", thoiGianDat=\"2026-03-26T09:15:00\""],
            ["SanPham", "sp1 : SanPham",
             "id=SP-0099, ten=\"Samsung Galaxy A56\", moTa=\"Điện thoại thông minh\", donGia=8.500.000 VND"],
            ["ChiTietDonHang", "ctd1 : ChiTietDonHang",
             "id=CTD-001-01, soLuong=1, donGia=8.500.000 VND"],
        ],
        "relations": [
            "kh1 đặt dh1  (KhachHang → DonHang, tên liên kết: đặt)",
            "dh1 gồm ctd1  (DonHang → ChiTietDonHang, tên liên kết: gồm)",
            "ctd1 tham chiếu sp1  (ChiTietDonHang → SanPham, tên liên kết: thuộc)",
        ],
    },
    "4.2": {
        "title": "Kịch bản: Hệ thống tự động gán tài xế cho đơn hàng mới",
        "narrative": (
            "Ngay sau khi ORD-2026-001 được tạo lúc 09:15, hệ thống phân công kích hoạt "
            "và quét 3 tài xế đang khả dụng trong bán kính 5km. Tài xế Trần Minh Khoa "
            "(TX-001) ở khoảng cách 2,1km, tải trọng hiện tại bằng 0 đơn và điểm hiệu suất "
            "4,8/5 — đứng đầu danh sách. Hệ thống tạo phiếu phân công PPC-2026-0326-001 "
            "và đẩy thông báo tới thiết bị của TX-001. Tài xế nhấn \"Chấp nhận\" lúc 09:16. "
            "Trạng thái DonHang chuyển sang Đã phân công."
        ),
        "table_headers": ["Lớp", "Tên đối tượng", "Giá trị thuộc tính chính"],
        "table_rows": [
            ["DonHang", "dh1 : DonHang",
             "id=ORD-2026-001, trangThai=\"Đã phân công\" (cập nhật từ UC-01)"],
            ["TaiXe", "tx1 : TaiXe",
             "id=TX-001, ten=\"Trần Minh Khoa\", soDienThoai=\"0912345678\", trangThai=\"Đang giao hàng\", viTriHienTai=\"10.7626, 106.6602\""],
            ["PhieuPhanCong", "ppc1 : PhieuPhanCong",
             "id=PPC-2026-0326-001, thoiGianPhanCong=\"2026-03-26T09:15:30\", trangThai=\"Đã chấp nhận\", diemUuTien=4.8"],
            ["VungDiaChi", "vda1 : VungDiaChi",
             "id=VDA-Q5-BCHB, tenVung=\"Q.5 – Bình Chánh\", banKinhTimKiem=5"],
            ["DiemHieuSuat", "dhs1 : DiemHieuSuat",
             "id=DHS-TX-001, diemTrungBinh=4.8, tongChuyenDi=142, tyLeHoanThanh=0.98"],
        ],
        "relations": [
            "ppc1 phân công dh1  (PhieuPhanCong → DonHang, tên liên kết: phân công)",
            "ppc1 gán tx1  (PhieuPhanCong → TaiXe, tên liên kết: gán)",
            "tx1 thuộc vùng vda1  (TaiXe → VungDiaChi, tên liên kết: thuộc vùng)",
            "tx1 có dhs1  (TaiXe → DiemHieuSuat, tên liên kết: có)",
        ],
    },
    "4.3": {
        "title": "Kịch bản: Tài xế lấy hàng tại kho và giao thành công đến khách",
        "narrative": (
            "09:30 — TX-001 đến kho KHO-SGN-01 (Kho Bình Chánh), quét mã QR trên kiện "
            "hàng KH-PKG-001; hệ thống xác nhận đúng đơn ORD-2026-001. Hệ thống tính lộ "
            "trình LT-2026-001 (9,2 km, ETA 35 phút) và chuyến đi CD-2026-0326-T1 bắt đầu. "
            "09:31 — tài xế xuất phát; hệ thống cập nhật vị trí GPS mỗi 30 giây. "
            "10:04 — TX-001 đến \"123 Nguyễn Văn Cừ, Q.5\"; hệ thống gửi thông báo "
            "\"Tài xế đang đến\" đến KH-055. "
            "10:06 — giao hàng thành công; TX-001 chụp ảnh biên nhận và ghi tên người "
            "nhận \"Nguyễn Thị Lan\" vào bằng chứng BC-2026-001. Trạng thái DonHang "
            "chuyển sang Giao thành công; sự kiện kích hoạt UC-04."
        ),
        "table_headers": ["Lớp", "Tên đối tượng", "Giá trị thuộc tính chính"],
        "table_rows": [
            ["TaiXe", "tx1 : TaiXe",
             "id=TX-001, ten=\"Trần Minh Khoa\", trangThai=\"Đã đến điểm giao\""],
            ["DonHang", "dh1 : DonHang",
             "id=ORD-2026-001, trangThai=\"Giao thành công\""],
            ["KienHang", "kh_pkg1 : KienHang",
             "id=KH-PKG-001, maQR=\"QR-ORD-2026-001\", khoiLuong=0.35 kg, trangThai=\"Đã giao\""],
            ["LoTrinh", "lt1 : LoTrinh",
             "id=LT-2026-001, diemBatDau=\"KHO-SGN-01\", diemDen=\"123 Nguyễn Văn Cừ, Q.5\", khoangCach=9.2 km, trangThai=\"Hoàn tất\""],
            ["ViTriGPS", "gps1 : ViTriGPS",
             "id=GPS-TX001-1006, viDo=10.7525, kinhDo=106.6624, thoiGian=\"2026-03-26T10:06:00\""],
            ["ChuyenDi", "cd1 : ChuyenDi",
             "id=CD-2026-0326-T1, thoiGianBatDau=\"2026-03-26T09:30:00\", thoiGianKetThuc=\"2026-03-26T10:10:00\""],
            ["BangChungGiaoHang", "bc1 : BangChungGiaoHang",
             "id=BC-2026-001, tenNguoiNhan=\"Nguyễn Thị Lan\", hinhAnh=\"proof_ORD-2026-001.jpg\", toaDo=\"10.7525, 106.6624\", thoiGianGiao=\"2026-03-26T10:06:00\""],
        ],
        "relations": [
            "cd1 bao gồm lt1  (ChuyenDi → LoTrinh, tên liên kết: bao gồm)",
            "tx1 thực hiện lt1  (TaiXe → LoTrinh, tên liên kết: thực hiện)",
            "lt1 giao dh1  (LoTrinh → DonHang, tên liên kết: giao)",
            "lt1 ghi nhận gps1  (LoTrinh → ViTriGPS, tên liên kết: ghi nhận)",
            "dh1 chứa kh_pkg1  (DonHang → KienHang, tên liên kết: chứa)",
            "dh1 có bc1  (DonHang → BangChungGiaoHang, tên liên kết: có bằng chứng)",
        ],
    },
    "4.4": {
        "title": "Kịch bản: Shipper xác nhận giao hàng bằng OTP",
        "narrative": (
            "10:07 — TX-001 (đóng vai Shipper) chọn ORD-2026-001 trong tab \"Chờ xác nhận\" "
            "trên ứng dụng rồi nhấn \"Xác nhận giao hàng\". Hệ thống gửi OTP 6 chữ số "
            "\"482913\" đến số điện thoại 0901234567 của KH-055 qua SMS. Khách hàng đọc mã "
            "cho shipper. TX-001 nhập mã vào ứng dụng; hệ thống kiểm tra hợp lệ. "
            "Hệ thống tạo bản ghi XacNhanGiaoHang XN-2026-001 (phương thức: OTP, "
            "thời gian: 10:07:44). Trạng thái DonHang chuyển sang Đã xác nhận. "
            "Hệ thống gửi thông báo TB-2026-001 đến KH-055 và Phòng Kế toán."
        ),
        "table_headers": ["Lớp", "Tên đối tượng", "Giá trị thuộc tính chính"],
        "table_rows": [
            ["TaiXe", "tx1 : TaiXe",
             "id=TX-001, ten=\"Trần Minh Khoa\", trangThai=\"Hoàn tất chuyến\""],
            ["DonHang", "dh1 : DonHang",
             "id=ORD-2026-001, trangThai=\"Đã xác nhận\""],
            ["KhachHang", "kh1 : KhachHang",
             "id=KH-055, ten=\"Nguyễn Thị Lan\", soDienThoai=\"0901234567\""],
            ["BangChungGiaoHang", "bc1 : BangChungGiaoHang",
             "id=BC-2026-001 (tham chiếu từ UC-03)"],
            ["XacNhanGiaoHang", "xn1 : XacNhanGiaoHang",
             "id=XN-2026-001, phuongThucXacNhan=\"OTP\", thoiGianXacNhan=\"2026-03-26T10:07:44\""],
            ["ThongBao", "tb1 : ThongBao",
             "id=TB-2026-001, noiDung=\"Đơn ORD-2026-001 đã giao thành công\", thoiGianGui=\"2026-03-26T10:07:45\""],
        ],
        "relations": [
            "tx1 lập xn1  (TaiXe → XacNhanGiaoHang, tên liên kết: lập)",
            "xn1 xác nhận dh1  (XacNhanGiaoHang → DonHang, tên liên kết: xác nhận)",
            "bc1 đính kèm xn1  (BangChungGiaoHang → XacNhanGiaoHang, tên liên kết: đính kèm)",
            "xn1 kích hoạt tb1  (XacNhanGiaoHang → ThongBao, tên liên kết: kích hoạt)",
            "kh1 nhận dh1  (KhachHang → DonHang, tên liên kết: nhận đơn)",
        ],
    },
    "4.5": {
        "title": "Kịch bản: Hệ thống xử lý COD sau khi giao hàng được xác nhận",
        "narrative": (
            "10:08 — hệ thống phát hiện ORD-2026-001 có hình thức thanh toán COD, tự động "
            "mở màn hình nhập số tiền thu thực cho TX-001. Shipper nhập 8.500.000 VND "
            "(khớp với tongGiaTri). Hệ thống tính phí dịch vụ PHI-2026-001: 85.000 VND "
            "(1% tổng giá trị) và hoa hồng tài xế: 42.500 VND (0,5%). "
            "Hệ thống tạo GiaoDich GD-2026-001 và HoaDon HD-2026-001. "
            "Cộng 42.500 VND vào ViDienTu VI-TX-001 của TX-001; ghi bút toán vào SoCai "
            "SC-2026-Q1. Trạng thái DonHang chuyển sang Hoàn tất thanh toán. "
            "Hệ thống phát sự kiện \"PaymentCompleted\"."
        ),
        "table_headers": ["Lớp", "Tên đối tượng", "Giá trị thuộc tính chính"],
        "table_rows": [
            ["TaiXe", "tx1 : TaiXe",
             "id=TX-001, ten=\"Trần Minh Khoa\""],
            ["DonHang", "dh1 : DonHang",
             "id=ORD-2026-001, trangThai=\"Hoàn tất thanh toán\", hinhThucThanhToan=\"COD\""],
            ["GiaoDich", "gd1 : GiaoDich",
             "id=GD-2026-001, soTienCOD=8.500.000 VND, phiDichVu=85.000 VND, hoaHong=42.500 VND, thoiGian=\"2026-03-26T10:08:00\", trangThai=\"Hoàn tất\""],
            ["HoaDon", "hd1 : HoaDon",
             "id=HD-2026-001, tongTien=8.500.000 VND, trangThai=\"Đã thanh toán\", thoiGianXuatHoaDon=\"2026-03-26T10:08:12\""],
            ["ViDienTu", "vdt1 : ViDienTu",
             "id=VDT-TX001, soDu=1.285.000 VND, chuSoHuu=\"TX-001\""],
            ["SoCai", "sc1 : SoCai",
             "id=SC-2026-Q1, kyKeToan=\"Q1-2026\", tongDoanhThu=85.000 VND"],
            ["PhiDichVu", "phi1 : PhiDichVu",
             "id=PHI-2026-001, tenLoaiPhi=\"Phí giao hàng nội thành\", tyLe=0.01, soTienApDung=85.000 VND"],
        ],
        "relations": [
            "dh1 kết toán gd1  (DonHang → GiaoDich, tên liên kết: kết toán)",
            "gd1 tạo hd1  (GiaoDich → HoaDon, tên liên kết: tạo)",
            "gd1 áp dụng phi1  (GiaoDich → PhiDichVu, tên liên kết: áp dụng)",
            "tx1 sở hữu vdt1  (TaiXe → ViDienTu, tên liên kết: sở hữu)",
            "gd1 cộng vào vdt1  (GiaoDich → ViDienTu, tên liên kết: cộng vào)",
            "gd1 ghi vào sc1  (GiaoDich → SoCai, tên liên kết: ghi vào)",
        ],
    },
}

# ─── Helpers ─────────────────────────────────────────────────────────────────

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def ft(para):
    return "".join(r.text for r in para.runs) or para.text

def insert_before(parent, new_elem, ref_elem):
    idx = list(parent).index(ref_elem)
    parent.insert(idx, new_elem)

def make_para(doc, text, style="Normal", bold_prefix=None):
    """Create a detached paragraph element."""
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_bold_label_para(doc, label, text):
    """Paragraph with a bold label followed by normal text."""
    p = doc.add_paragraph(style="Normal")
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(text)
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_heading4(doc, text):
    # Use Heading 4 style if exists, else Normal+bold
    try:
        p = doc.add_paragraph(text, style="Heading 4")
    except Exception:
        p = doc.add_paragraph(style="Normal")
        r = p.add_run(text)
        r.bold = True
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_table(doc, headers, rows):
    """Create a table and detach it from the document body, returning (tbl_elem, sep_elem)."""
    num_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
    tbl.style = "TableNormal"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            row_cells[ci].text = val

    tbl_elem = tbl._tbl
    tbl_elem.getparent().remove(tbl_elem)

    # Also make a blank paragraph to follow the table
    spacer = doc.add_paragraph()
    spacer_elem = spacer._element
    spacer_elem.getparent().remove(spacer_elem)

    return tbl_elem, spacer_elem

def make_bullet(doc, text):
    """Bullet-style paragraph (List Bullet if available, else Normal with –)."""
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except Exception:
        p = doc.add_paragraph("– " + text, style="Normal")
    elem = p._element
    elem.getparent().remove(elem)
    return elem

# ─── Main transformation ──────────────────────────────────────────────────────

doc = Document(DOC_PATH)
body = doc.element.body

import re
UC_RE = re.compile(r"^(4\.\d+)\.")

# Find scenario placeholder paragraphs: "[Mô tả kịch bản..."
# The placeholder follows a Heading 3 containing "Kịch bản"
# We replace the placeholder element in-place by inserting new elements before it and removing it.

to_process = []  # list of (uc_num, placeholder_elem)

paras = doc.paragraphs
for i, p in enumerate(paras):
    t = ft(p)
    if "Mô tả kịch bản cụ thể" in t or "bảng đối tượng" in t:
        # Find which UC this belongs to by scanning back for the nearest Heading 2
        uc_num = None
        for j in range(i - 1, -1, -1):
            m = UC_RE.match(ft(paras[j]))
            if m and paras[j].style.name.startswith("Heading 2"):
                uc_num = m.group(1)
                break
        if uc_num:
            to_process.append((uc_num, p._element))

print(f"Found {len(to_process)} scenario placeholders: {[x[0] for x in to_process]}")

for uc_num, placeholder_elem in to_process:
    data = UC_SCENARIOS.get(uc_num)
    if not data:
        print(f"  No data for {uc_num}, skipping")
        continue

    parent = placeholder_elem.getparent()
    ref = placeholder_elem  # we'll insert before this, then remove it

    new_elems = []

    # 1. Title line: bold "Tên kịch bản: " + title text
    new_elems.append(make_bold_label_para(doc, "Tên kịch bản: ", data["title"]))

    # 2. Blank spacer
    new_elems.append(make_para(doc, ""))

    # 3. Narrative
    new_elems.append(make_para(doc, data["narrative"]))

    # 4. Blank spacer
    new_elems.append(make_para(doc, ""))

    # 5. Object table heading
    new_elems.append(make_bold_label_para(doc, "Bảng đối tượng:", ""))

    # 6. Table
    tbl_elem, spacer_elem = make_table(doc, data["table_headers"], data["table_rows"])
    new_elems.append(tbl_elem)
    new_elems.append(spacer_elem)

    # 7. Relations heading
    new_elems.append(make_bold_label_para(doc, "Quan hệ đối tượng:", ""))

    # 8. Bullet list
    for rel in data["relations"]:
        new_elems.append(make_bullet(doc, rel))

    # Insert all before placeholder, then remove placeholder
    for elem in new_elems:
        insert_before(parent, elem, ref)
    parent.remove(ref)

    print(f"  Filled {uc_num} ({len(new_elems)} elements inserted)")

# ─── Save ────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Done. Saved to", DOC_PATH)
